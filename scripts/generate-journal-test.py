#!/usr/bin/env python3
"""期刊论文测试：对齐《煤炭工程》类双栏格式，基于用户提供的 PDF 样例结构。"""
from __future__ import annotations

import importlib.util
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent.parent
SCRIPT = ROOT / "scripts" / "word-typeset.py"
DATA = ROOT / "scripts" / "word-typeset-test-data"
ASSETS = DATA / "assets"
TEXT_SRC = DATA / "期刊论文-测试原文.txt"
DOC_IN = DATA / "期刊论文-排版前.docx"
DOC_OUT = DATA / "期刊论文-排版后.docx"
FLOW_PNG = ASSETS / "journal-gangue-flow.png"


def load_module():
    spec = importlib.util.spec_from_file_location("word_typeset", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


def find_cjk_font(size: int = 16):
    for path in (
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "C:/Windows/Fonts/msyh.ttc",
    ):
        if Path(path).is_file():
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                continue
    return ImageFont.load_default()


def draw_gangue_flow_chart(path: Path) -> None:
    w, h = 680, 300
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = find_cjk_font(20)
    box_font = find_cjk_font(16)
    draw.text((36, 20), "煤矸石综合利用技术路线", fill=(30, 64, 175), font=title_font)
    boxes = [
        (40, 90, 180, 160, "破碎分级", (219, 234, 254), (59, 130, 246)),
        (250, 90, 390, 160, "充填开采", (220, 252, 231), (34, 197, 94)),
        (460, 90, 600, 160, "建材制备", (254, 243, 199), (245, 158, 11)),
    ]
    for x0, y0, x1, y1, label, fill, outline in boxes:
        draw.rounded_rectangle((x0, y0, x1, y1), radius=8, fill=fill, outline=outline, width=2)
        bbox = draw.textbbox((0, 0), label, font=box_font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text((x0 + (x1 - x0 - tw) / 2, y0 + (y1 - y0 - th) / 2), label, fill=(33, 37, 41), font=box_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, format="PNG")


def build_docx_from_text(text: str, path: Path) -> None:
    draw_gangue_flow_chart(FLOW_PNG)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    doc = Document()
    if not lines:
        raise ValueError("原文为空")

    doc.add_paragraph(lines[0], style="Title")
    if len(lines) > 1:
        doc.add_paragraph(lines[1], style="Subtitle")

    i = 2
    while i < len(lines):
        line = lines[i]
        if line.startswith("（") or line.startswith("("):
            doc.add_paragraph(line)
            i += 1
            continue
        if line.startswith("表") and re.match(r"^表\s*[\d一二三四五六七八九十]+", line):
            doc.add_paragraph(line)
            i += 1
            if i < len(lines) and re.match(r"^Table\s+\d+", lines[i], re.I):
                doc.add_paragraph(lines[i])
                i += 1
            table = doc.add_table(rows=4, cols=4)
            hdr = ["来源", "表观密度/(kg·m⁻³)", "堆积密度/(kg·m⁻³)", "备注"]
            for c, h in enumerate(hdr):
                table.rows[0].cells[c].text = h
            rows = [
                ("山西 A 矿", "2100～2400", "1300～1500", "高碳"),
                ("内蒙古 B 矿", "2200～2500", "1400～1600", "高铁"),
                ("陕西 C 矿", "2000～2300", "1250～1450", "风化"),
            ]
            for r, row in enumerate(rows, 1):
                for c, val in enumerate(row):
                    table.rows[r].cells[c].text = val
            continue
        if line.startswith("图"):
            pic = doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.add_run().add_picture(str(FLOW_PNG), width=Inches(4.6))
            doc.add_paragraph(line)
            i += 1
            continue
        if re.match(r"^\d{1,2}\s+[\u4e00-\u9fff]", line):
            doc.add_paragraph(line)
        elif re.match(r"^\d+\.\s*\d+", line):
            doc.add_paragraph(line)
        elif line.startswith(("摘要", "关键词", "中图", "Abstract", "Keywords", "收稿", "作者", "引用", "doi")):
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line, style="Normal")
        i += 1

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def run_typeset(inp: Path, out: Path, config: dict) -> None:
    py = ROOT / ".venv" / "bin" / "python3"
    if not py.is_file():
        py = Path(sys.executable)
    payload = {
        "mode": "batch",
        "input_paths": [str(inp)],
        "output_path": str(out),
        "in_place": False,
        "config": config,
    }
    proc = subprocess.run(
        [str(py), str(SCRIPT), json.dumps(payload, ensure_ascii=False)],
        capture_output=True,
        text=True,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr or proc.stdout)
    result = json.loads(proc.stdout.strip())
    if not result.get("ok"):
        raise RuntimeError(result.get("error", "排版失败"))


def section_columns(doc: Document) -> list[int]:
    counts = []
    for section in doc.sections:
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is not None and cols.get(qn("w:num")):
            counts.append(int(cols.get(qn("w:num"))))
        else:
            counts.append(1)
    return counts


def header_text(doc: Document) -> str:
    parts: list[str] = []
    for section in doc.sections:
        for header in (section.header, section.first_page_header):
            try:
                text = "\n".join(p.text for p in header.paragraphs if p.text.strip())
            except Exception:
                text = ""
            if text.strip():
                parts.append(text.strip())
    return " | ".join(parts)


def verify(doc_path: Path, config: dict) -> list[tuple[str, bool, str]]:
    doc = Document(str(doc_path))
    checks: list[tuple[str, bool, str]] = []

    cols = section_columns(doc)
    checks.append(("分节双栏", len(cols) >= 2 and cols[-1] == 2, str(cols)))

    intro = next((p for p in doc.paragraphs if p.text.startswith("煤炭是我国")), None)
    checks.append(("引言在双栏节", intro is not None, "找到" if intro else "未找到"))

    h1 = next((p for p in doc.paragraphs if p.text.startswith("1 选煤")), None)
    if h1 and h1.runs:
        r = h1.runs[0]._element.rPr
        fonts = r.find(qn("w:rFonts")) if r is not None else None
        east = fonts.get(qn("w:eastAsia")) if fonts is not None else h1.runs[0].font.name
        checks.append(("一级标题黑体", "黑" in (east or ""), east or ""))
    else:
        checks.append(("一级标题黑体", False, "未找到"))

    abstract = next((p for p in doc.paragraphs if p.text.startswith("摘要")), None)
    if abstract:
        pf = abstract.paragraph_format
        hang_ok = pf.left_indent and pf.first_line_indent and pf.first_line_indent < 0
        checks.append(("摘要悬挂缩进", bool(hang_ok), f"左{pf.left_indent} 首{pf.first_line_indent}"))
    else:
        checks.append(("摘要悬挂缩进", False, "未找到"))

    table_en = next((p for p in doc.paragraphs if p.text.startswith("Table 1")), None)
    checks.append(("英文表题", table_en is not None, table_en.text[:40] if table_en else "未找到"))

    cite_para = next((p for p in doc.paragraphs if "[1]" in p.text), None)
    if cite_para:
        sup = any(r.font.superscript for r in cite_para.runs if r.text)
        checks.append(("引文上标", sup, f"runs={len(cite_para.runs)}"))
    else:
        checks.append(("引文上标", False, "未找到"))

    hdr = header_text(doc)
    checks.append(
        (
            "期刊页眉",
            "煤炭工程" in hdr.replace(" ", "") and "专家论坛" in hdr,
            hdr[:80] or "空",
        )
    )

    checks.append(("表格保留", len(doc.tables) >= 1, f"{len(doc.tables)} 张"))
    inline = sum(len(p._element.xpath(".//wp:inline")) for p in doc.paragraphs)
    checks.append(("插图保留", inline >= 1, f"{inline} 个"))

    left = round(doc.sections[0].left_margin.cm, 2)
    checks.append(
        (
            "左边距约1.8cm",
            abs(left - config["page"]["margin_left"]) < 0.25,
            f"{left} cm",
        )
    )
    return checks


def main() -> int:
    wt = load_module()
    if not hasattr(wt, "journal_default_config"):
        raise RuntimeError("word-typeset.py 缺少 journal_default_config()")
    text = TEXT_SRC.read_text(encoding="utf-8")
    print("==> 期刊论文测试（对齐煤炭工程类格式）")
    build_docx_from_text(text, DOC_IN)
    print(f"    生成: {DOC_IN}")
    if DOC_OUT.exists():
        DOC_OUT.unlink()
    config = wt.journal_default_config()
    run_typeset(DOC_IN, DOC_OUT, config)
    print(f"    排版: {DOC_OUT}")
    checks = verify(DOC_OUT, config)
    failed = 0
    print("\n==> 期刊格式校验")
    for name, ok, detail in checks:
        mark = "✓" if ok else "✗"
        print(f"    {mark} {name}: {detail}")
        if not ok:
            failed += 1
    if failed:
        return 1
    print("\n完成。可用 Word 打开「期刊论文-排版后.docx」与 PDF 样例对照。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
