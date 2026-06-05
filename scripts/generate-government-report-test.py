#!/usr/bin/env python3
"""生成约1000字政府工作报告测试稿，并执行排版。"""
from __future__ import annotations

import importlib.util
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
SCRIPT = ROOT / "scripts" / "word-typeset.py"
DATA = ROOT / "scripts" / "word-typeset-test-data"
TEXT_SRC = DATA / "政府工作报告-测试原文.txt"
DOC_IN = DATA / "政府工作报告-排版前.docx"
DOC_OUT = DATA / "政府工作报告-排版后.docx"


def load_module():
    spec = importlib.util.spec_from_file_location("word_typeset", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


def char_count(text: str) -> int:
    return len("".join(ch for ch in text if not ch.isspace()))


def build_docx_from_text(text: str, path: Path) -> None:
    """按机关公文结构生成未排版 docx。"""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    doc = Document()
    if not lines:
        raise ValueError("原文为空")
    doc.add_paragraph(lines[0], style="Title")
    if len(lines) > 1:
        doc.add_paragraph(lines[1], style="Subtitle")
    i = 2
    while i < len(lines):
        line = lines[i]
        if line.startswith("表") and re.match(r"^表\s*[\d一二三四五六七八九十]+", line):
            doc.add_paragraph(line)
            table = doc.add_table(rows=4, cols=4)
            hdr = ["指标名称", "年度目标", "实际完成", "完成率"]
            for c, h in enumerate(hdr):
                table.rows[0].cells[c].text = h
            rows_data = [
                ("地区生产总值增速", "5.5%", "5.8%", "105%"),
                ("固定资产投资增速", "6.5%", "7.1%", "109%"),
                ("城镇新增就业", "8万人", "8.6万人", "108%"),
            ]
            for r, row in enumerate(rows_data, 1):
                for c, val in enumerate(row):
                    table.rows[r].cells[c].text = val
            i += 1
            continue
        if line.startswith("图"):
            doc.add_paragraph(line)
            i += 1
            continue
        if line.startswith("附件"):
            doc.add_paragraph(line)
            i += 1
            continue
        if line.startswith("（"):
            doc.add_paragraph(line, style="Heading 2")
        elif line.startswith(("一、", "二、", "三、", "四、", "五、")):
            doc.add_paragraph(line, style="Heading 1")
        else:
            doc.add_paragraph(line, style="Normal")
        i += 1
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def run_typeset(inp: Path, out: Path, config: dict) -> None:
    py = ROOT / ".venv" / "bin" / "python3"
    if not py.is_file():
        py = Path(sys.executable)
    payload = {
        "mode": "batch",
        "input_paths": [str(inp)],
        "output_path": str(out),
        "in_place": False,
        "config": config,
    }
    proc = subprocess.run(
        [str(py), str(SCRIPT), json.dumps(payload, ensure_ascii=False)],
        capture_output=True,
        text=True,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr or proc.stdout)
    result = json.loads(proc.stdout.strip())
    if not result.get("ok"):
        raise RuntimeError(result.get("error", "排版失败"))


def main() -> int:
    wt = load_module()
    text = TEXT_SRC.read_text(encoding="utf-8")
    n = char_count(text)
    print(f"==> 政府工作报告测试稿（汉字约 {n} 字）")
    print(f"    原文: {TEXT_SRC}")

    build_docx_from_text(text, DOC_IN)
    print(f"    生成: {DOC_IN}")

    if DOC_OUT.exists():
        DOC_OUT.unlink()
    config = wt.builtin_default_config()
    run_typeset(DOC_IN, DOC_OUT, config)
    print(f"    排版: {DOC_OUT}")

    doc = Document(str(DOC_OUT))
    print(f"    段落 {len(doc.paragraphs)} 段，表格 {len(doc.tables)} 张")
    print("完成。可用 Word 打开「政府工作报告-排版后.docx」查看效果。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
