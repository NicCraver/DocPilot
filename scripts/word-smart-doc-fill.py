#!/usr/bin/env python3
"""按 profile 把内容灌入 original.docx 骨架，产出排版后的 docx。"""
from __future__ import annotations

import json
import re
import shutil
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
GUIDE_KEYWORDS = ("填写指南", "请点击", "扫描二维码", "feishu.cn", "格式刷", "使用时")


def _norm(t: str) -> str:
    return re.sub(r"\s+", "", t.strip())


def _run_is_red(run) -> bool:
    color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
    return color == RGBColor(0xFF, 0x00, 0x00)


def paragraph_is_guide(para: Paragraph) -> bool:
    text = para.text.strip()
    if not text:
        return False
    red = sum(1 for r in para.runs if _run_is_red(r))
    if red and (text.startswith("（") or red >= max(1, len(para.runs) // 2)):
        return True
    if any(k in text for k in GUIDE_KEYWORDS):
        return True
    return False


def is_placeholder(para: Paragraph) -> bool:
    text = para.text.strip()
    return text in {"（此处填写正文）", "（此处填写内容）"} or paragraph_is_guide(para)


def set_east_asia(run, font: str) -> None:
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts if rpr.rFonts is not None else OxmlElement("w:rFonts")
    if rpr.rFonts is None:
        rpr.append(rfonts)
    rfonts.set(f"{NS}eastAsia", font)


def apply_body_style(para: Paragraph, style: dict[str, Any]) -> None:
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = style.get("line_spacing", 1.5)
    indent_chars = style.get("first_line_indent_chars", 0)
    size_pt = style.get("size_pt", 12)
    if indent_chars:
        pf.first_line_indent = Pt(size_pt * indent_chars)
    for run in para.runs:
        if style.get("font_ea"):
            set_east_asia(run, style["font_ea"])
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.italic = False
        run.bold = bool(style.get("bold", False))


def insert_after(anchor: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        para.add_run(text)
    return para


def remove_para(para: Paragraph) -> None:
    el = para._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def split_paragraphs(text: str) -> list[str]:
    out = []
    for chunk in re.split(r"\n\s*\n+", text.replace("\r\n", "\n")):
        line = chunk.strip()
        if line:
            out.append(line)
    return out


def parse_markdown(raw: str, structure: list[dict]) -> tuple[dict[str, list[str]], dict[str, str]]:
    titles = {_norm(s["title"]): s["key"] for s in structure}
    sections: dict[str, list[str]] = {s["key"]: [] for s in structure}
    meta: dict[str, str] = {}
    order = [s["key"] for s in structure]
    cur = order[0] if order else None
    buf: list[str] = []

    def flush():
        nonlocal buf
        if cur and buf:
            sections.setdefault(cur, []).extend(split_paragraphs("\n".join(buf)))
        buf = []

    for line in raw.replace("\r\n", "\n").split("\n"):
        s = line.strip()
        if not s:
            buf.append("")
            continue
        m = re.match(r"^#{1,6}\s+(.+)$", s)
        head = m.group(1).strip() if m else s
        mr = re.match(r"汇报人[：:]\s*(.+)$", s)
        md = re.match(r"日期[：:]\s*(.+)$", s)
        if mr:
            meta["reporter"] = mr.group(1).strip(); continue
        if md:
            meta["date"] = md.group(1).strip(); continue
        key = titles.get(_norm(head))
        if key and (m or len(s) < 40):
            flush(); cur = key; continue
        if cur:
            buf.append(s)
    flush()
    return sections, meta


def load_sections(payload: dict, structure: list[dict]) -> tuple[dict[str, list[str]], dict[str, str]]:
    kind = payload.get("content_kind", "text")
    if kind == "sections" or payload.get("sections"):
        sections = payload.get("sections")
        if not sections:
            raise ValueError("content_kind=sections 但未提供 sections")
        return {k: list(v) for k, v in sections.items()}, {}
    text = payload.get("content_text")
    if not text and payload.get("content_path"):
        p = Path(payload["content_path"])
        if p.suffix.lower() == ".docx":
            doc = Document(str(p))
            lines = [para.text.strip() for para in doc.paragraphs
                     if para.text.strip() and not paragraph_is_guide(para)]
            text = "\n".join(lines)
        else:
            text = p.read_text(encoding="utf-8")
    if not text:
        raise ValueError("请提供内容文件、文本或 sections")
    return parse_markdown(text, structure)


def fill(payload: dict) -> dict[str, Any]:
    logs: list[str] = []
    tdir = Path(payload["template_dir"])
    original = tdir / "original.docx"
    profile = json.loads((tdir / "profile.json").read_text(encoding="utf-8"))
    structure = profile.get("structure", [])
    styles = profile.get("styles", {})
    body_style = styles.get("body", {"size_pt": 12, "first_line_indent_chars": 2, "line_spacing": 1.5})

    out = Path(payload["output_path"])
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(original, out)
    doc = Document(str(out))

    sections, meta = load_sections(payload, structure)
    reporter = payload.get("reporter") or meta.get("reporter")
    report_date = payload.get("report_date") or meta.get("date")

    removed = 0
    for para in list(doc.paragraphs):
        if is_placeholder(para):
            remove_para(para); removed += 1
    logs.append(f"已移除 {removed} 段占位/说明")

    titles = {_norm(s["title"]): s["key"] for s in structure}
    filled: set[str] = set()
    for para in list(doc.paragraphs):
        key = titles.get(_norm(para.text))
        if not key:
            continue
        parts = sections.get(key, [])
        if not parts:
            continue
        anchor = para
        for part in parts:
            np = insert_after(anchor, part)
            apply_body_style(np, body_style)
            anchor = np
        filled.add(key)
        logs.append(f"已填入章节「{para.text.strip()[:30]}」{len(parts)} 段")

    leftover = [(s["key"], sections.get(s["key"], [])) for s in structure
                if s["key"] not in filled and sections.get(s["key"])]
    if leftover:
        for key, parts in leftover:
            doc.add_paragraph(next((s["title"] for s in structure if s["key"] == key), key))
            for part in parts:
                p = doc.add_paragraph(part)
                apply_body_style(p, body_style)
        logs.append(f"{len(leftover)} 个未匹配章节已续写到文末")

    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t.startswith("汇报人") and reporter:
            for r in list(para.runs):
                r._r.getparent().remove(r._r)
            para.add_run(f"汇报人：{reporter}")
            logs.append(f"已填写汇报人：{reporter}")
        elif t.startswith("日期") and report_date:
            for r in list(para.runs):
                r._r.getparent().remove(r._r)
            para.add_run(f"日期：{report_date}")
            logs.append(f"已填写日期：{report_date}")

    doc.save(str(out))
    logs.append(f"已生成: {out}")
    return {"ok": True, "results": [{"input": payload.get("content_path") or "(text/sections)", "output": str(out)}], "logs": logs}


def main() -> None:
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "缺少 JSON 参数"}, ensure_ascii=False)); sys.exit(1)
    try:
        payload = json.loads(sys.argv[1])
        print(json.dumps(fill(payload), ensure_ascii=False))
    except Exception as e:  # noqa: BLE001
        print(json.dumps({"ok": False, "error": str(e)}, ensure_ascii=False)); sys.exit(1)


if __name__ == "__main__":
    main()
