#!/usr/bin/env python3
"""
Word 批量排版自动化测试（多场景）：
- 基础 / 复杂 / 启发式标题 / 自定义配置
- 批量多文件 / 原地备份 / 表格禁用
"""
from __future__ import annotations

import importlib.util
import json
import shutil
import subprocess
import sys
import tempfile
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
SCRIPT = ROOT / "scripts" / "word-typeset.py"
FIXTURES_DIR = ROOT / "scripts" / "word-typeset-test-data"
ARTIFACTS_DIR = ROOT / "scripts" / "test-artifacts" / "word-typeset"

# 测试 docx 统一使用中文文件名
F = {
    "基础_前": "基础样例-排版前.docx",
    "基础_后": "基础样例-排版后.docx",
    "复杂_前": "复杂样例-排版前.docx",
    "复杂_后": "复杂样例-排版后.docx",
    "自定义_后": "自定义配置-排版后.docx",
    "启发_前": "启发式标题-排版前.docx",
    "启发_后": "启发式标题-排版后.docx",
    "表格禁_原": "表格禁用-原始.docx",
    "表格禁_后": "表格禁用-排版后.docx",
    "原地_稿": "原地备份-工作稿.docx",
    "长文_前": "公文长文-排版前.docx",
    "长文_后": "公文长文-排版后.docx",
    "文本_出": "文本模式-输出.docx",
    "批量_A": "年度报告.docx",
    "批量_B": "工作方案.docx",
    "中文路径_文件": "《重点项目可行性研究报告》.docx",
    "中文路径_后": "《重点项目可行性研究报告》-排版后.docx",
}
DIR_BATCH = "批量排版"
DIR_CN_PATH = Path("中文路径") / "某市发改委" / "报送材料"

# 迁移前旧英文名（测试启动时清理）
LEGACY_FILES = [
    "sample-unformatted.docx",
    "sample-formatted.docx",
    "sample-complex-unformatted.docx",
    "sample-complex-formatted.docx",
    "sample-custom-formatted.docx",
    "sample-heuristic-unformatted.docx",
    "sample-heuristic-formatted.docx",
    "sample-table-raw.docx",
    "sample-table-disabled.docx",
    "inplace-work.docx",
    "inplace-work.docx.bak",
]

FONT_ALIASES: dict[str, set[str]] = {
    "方正小标宋简体": {"方正小标宋简体", "方正小标宋", "FZXiaoBiaoSong-B05", "FZXBSJW"},
    "宋体": {"宋体", "SimSun", "simsun", "Songti SC", "STSong"},
    "黑体": {"黑体", "SimHei", "simhei", "Heiti SC", "STHeiti"},
    "楷体": {"楷体", "KaiTi", "kaiti", "STKaiti"},
    "楷体_GB2312": {"楷体_GB2312", "楷体", "KaiTi", "KaiTi_GB2312", "STKaiti"},
    "仿宋_GB2312": {"仿宋_GB2312", "仿宋", "FangSong", "FangSong_GB2312", "STFangsong"},
    "Times New Roman": {"Times New Roman", "times new roman", "TimesNewRoman"},
    "Arial": {"Arial", "arial"},
}


def load_word_typeset_module():
    spec = importlib.util.spec_from_file_location("word_typeset", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


@dataclass
class CheckResult:
    name: str
    ok: bool
    detail: str = ""
    scenario: str = ""


@dataclass
class TestReport:
    checks: list[CheckResult] = field(default_factory=list)

    def add(self, scenario: str, name: str, ok: bool, detail: str = "") -> None:
        self.checks.append(CheckResult(name, ok, detail, scenario))

    @property
    def passed(self) -> bool:
        return all(c.ok for c in self.checks)

    def dump(self) -> dict[str, Any]:
        failed = [c for c in self.checks if not c.ok]
        by_scenario: dict[str, dict[str, Any]] = {}
        for c in self.checks:
            by_scenario.setdefault(c.scenario, {"passed": 0, "total": 0, "failed": []})
            by_scenario[c.scenario]["total"] += 1
            if c.ok:
                by_scenario[c.scenario]["passed"] += 1
            else:
                by_scenario[c.scenario]["failed"].append(c.name)
        return {
            "passed": self.passed,
            "total": len(self.checks),
            "failed": [f"{c.scenario}:{c.name}" for c in failed],
            "scenarios": by_scenario,
            "checks": [
                {
                    "scenario": c.scenario,
                    "name": c.name,
                    "ok": c.ok,
                    "detail": c.detail,
                }
                for c in self.checks
            ],
        }


def nearly(a: float, b: float, tol: float) -> bool:
    return abs(a - b) <= tol


def pt_value(length) -> float | None:
    if length is None:
        return None
    return float(length.pt)


def cm_value(length) -> float | None:
    if length is None:
        return None
    return float(length.cm)


def font_matches(actual: str | None, expected: str) -> bool:
    act = (actual or "").strip()
    if not act:
        return False
    aliases = FONT_ALIASES.get(expected, {expected})
    return act in aliases or act.lower() in {a.lower() for a in aliases}


def run_east_asia_font(run) -> str | None:
    r_pr = run._element.rPr
    if r_pr is None:
        return run.font.name
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        return run.font.name
    return r_fonts.get(qn("w:eastAsia")) or run.font.name


def run_ascii_font(run) -> str | None:
    r_pr = run._element.rPr
    if r_pr is None:
        return run.font.name
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        return run.font.name
    return r_fonts.get(qn("w:ascii")) or run.font.name


def paragraph_line_spacing_pt(paragraph) -> float | None:
    pf = paragraph.paragraph_format
    if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY:
        return None
    return pt_value(pf.line_spacing)


def footer_page_run_props(section) -> tuple[str | None, float | None]:
    footer = section.footer
    for p in footer.paragraphs:
        for run in p.runs:
            if run.font.size:
                return run_east_asia_font(run), pt_value(run.font.size)
    return None, None


def footer_has_page_field(section) -> bool:
    return "fldChar" in section.footer._element.xml and "PAGE" in section.footer._element.xml


def table_border_size_pt(table) -> float | None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return None
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        return None
    top = borders.find(qn("w:top"))
    if top is None:
        return None
    sz = top.get(qn("w:sz"))
    return int(sz) / 8.0 if sz else None


def outline_level(paragraph) -> int | None:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return None
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    val = outline.get(qn("w:val"))
    return int(val) if val is not None else None


def first_run_props(paragraph) -> tuple[str | None, float | None, bool | None]:
    if not paragraph.runs:
        return None, None, None
    run = paragraph.runs[0]
    return run_east_asia_font(run), pt_value(run.font.size), run.font.bold


def all_runs_east_asia_font(paragraph, expected: str) -> bool:
    if not paragraph.runs:
        return False
    return all(font_matches(run_east_asia_font(r), expected) for r in paragraph.runs)


def python_bin() -> Path:
    py = ROOT / ".venv" / "bin" / "python3"
    return py if py.is_file() else Path(sys.executable)


def run_typeset_payload(payload: dict[str, Any]) -> dict[str, Any]:
    proc = subprocess.run(
        [str(python_bin()), str(SCRIPT), json.dumps(payload, ensure_ascii=False)],
        capture_output=True,
        text=True,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr or proc.stdout or "排版脚本失败")
    result = json.loads(proc.stdout.strip())
    if not result.get("ok"):
        raise RuntimeError(result.get("error", "排版失败"))
    return result


def run_typeset_file(
    input_path: Path,
    output_path: Path,
    config: dict[str, Any],
    *,
    in_place: bool = False,
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "mode": "batch",
        "input_paths": [str(input_path)],
        "config": config,
        "in_place": in_place,
    }
    if not in_place:
        payload["output_path"] = str(output_path)
    return run_typeset_payload(payload)


def find_paragraph(doc: Document, predicate) -> Any | None:
    for p in doc.paragraphs:
        if predicate(p):
            return p
    return None


# ---------------------------------------------------------------------------
# 测试样例生成
# ---------------------------------------------------------------------------


def build_basic_fixture(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("2024年度工作总结报告", style="Title")
    doc.add_paragraph("某某单位", style="Subtitle")
    doc.add_paragraph("一、总体情况", style="Heading 1")
    doc.add_paragraph("（一）主要指标", style="Heading 2")
    doc.add_paragraph(
        "本文正文段落，包含中文与 English 123 混合内容，用于验证 ascii 字体与正文行距。",
        style="Normal",
    )
    doc.add_paragraph("表1 主要数据统计")
    table = doc.add_table(rows=3, cols=3)
    for i, h in enumerate(["指标", "数值", "单位"]):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text, table.rows[1].cells[1].text, table.rows[1].cells[2].text = (
        "产值",
        "100",
        "万元",
    )
    table.rows[2].cells[0].text, table.rows[2].cells[1].text, table.rows[2].cells[2].text = (
        "人数",
        "50",
        "人",
    )
    doc.add_paragraph("图1 组织架构示意图")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_complex_fixture(path: Path) -> None:
    """复杂样例：多表、混排 run、三级标题、附件、多空格题注。"""
    doc = Document()
    doc.add_paragraph("2025年度重点项目实施方案", style="Title")
    doc.add_paragraph("技术研发中心", style="Subtitle")
    doc.add_paragraph("一、项目背景", style="Heading 1")
    doc.add_paragraph("（一）政策依据", style="Heading 2")
    doc.add_paragraph("贯彻落实相关文件精神，推进数字化转型。", style="Normal")

    p_mix = doc.add_paragraph()
    r1 = p_mix.add_run("混排测试：")
    r1.bold = True
    p_mix.add_run("中文正文 English 789 与数字 3.14。")

    doc.add_paragraph("二、实施路径")  # 启发式一级（无样式）
    doc.add_paragraph("（二）阶段安排", style="Heading 2")
    doc.add_paragraph("1. 第一阶段完成需求梳理。", style="Normal")
    doc.add_paragraph("任务分解与里程碑", style="Heading 3")

    doc.add_paragraph("表1 里程碑计划")
    t1 = doc.add_table(rows=4, cols=4)
    t1.rows[0].cells[0].text = "阶段"
    t1.rows[0].cells[1].text = "时间"
    t1.rows[0].cells[2].text = "负责人"
    t1.rows[0].cells[3].text = "状态"
    t1.rows[1].cells[0].text = "启动"
    t1.rows[1].cells[1].text = "2025-Q1"
    t1.rows[1].cells[2].text = "张三"
    t1.rows[1].cells[3].text = "完成"
    t1.rows[2].cells[0].text = "开发"
    t1.rows[2].cells[1].text = "2025-Q2"
    t1.rows[2].cells[2].text = "李四"
    t1.rows[2].cells[3].text = "进行"
    t1.rows[3].cells[0].text = "验收"
    t1.rows[3].cells[1].text = "2025-Q3"
    t1.rows[3].cells[2].text = "王五"
    t1.rows[3].cells[3].text = "待定"

    doc.add_paragraph("图1 总体技术架构")
    doc.add_paragraph("表2 经费预算明细")
    t2 = doc.add_table(rows=3, cols=5)
    headers = ["科目", "预算(万元)", "已用", "余额", "备注"]
    for i, h in enumerate(headers):
        t2.rows[0].cells[i].text = h
    t2.rows[1].cells[0].text = "设备"
    t2.rows[1].cells[1].text = "120.5"
    t2.rows[1].cells[2].text = "80"
    t2.rows[1].cells[3].text = "40.5"
    t2.rows[1].cells[4].text = "含服务器"
    t2.rows[2].cells[0].text = "人力"
    t2.rows[2].cells[1].text = "200"
    t2.rows[2].cells[2].text = "150"
    t2.rows[2].cells[3].text = "50"
    t2.rows[2].cells[4].text = "研发"

    doc.add_paragraph("图 2 部署流程图")  # 空格题注
    doc.add_paragraph("附件1 申报材料清单")
    doc.add_paragraph("附件 2 合同扫描件")  # 空格附件编号

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_heuristic_fixture(path: Path) -> None:
    """纯文本启发式：无 Word 内置样式，仅靠序号规则识别标题。"""
    doc = Document()
    doc.add_paragraph("三、纯文本标题识别测试")
    doc.add_paragraph("（四）二级标题启发式")
    doc.add_paragraph("本段为普通正文，无内置 Heading 样式，仅通过排版引擎规则处理。")
    doc.add_paragraph("表3 启发式表题")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "列A"
    table.rows[0].cells[1].text = "列B"
    table.rows[1].cells[0].text = "数据1"
    table.rows[1].cells[1].text = "数据2"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_long_document_fixture(path: Path) -> None:
    """公文长文：多章多节、多表、多级标题。"""
    doc = Document()
    doc.add_paragraph("某某市人民政府公文格式测试稿", style="Title")
    doc.add_paragraph("办公室", style="Subtitle")

    sections = [
        ("一、总体要求", "（一）指导思想", "坚持高质量发展，统筹发展与安全。"),
        ("二、重点任务", "（二）工作举措", "细化任务清单，明确责任分工。"),
        ("三、保障措施", "（三）组织保障", "健全工作机制，强化督促检查。"),
        ("四、经费预算", None, "科学编制预算，严格财经纪律。"),
        ("五、实施计划", "（四）进度安排", "按季度推进，确保如期完成。"),
    ]
    for i, (h1, h2, body) in enumerate(sections, 1):
        doc.add_paragraph(h1, style="Heading 1")
        if h2:
            doc.add_paragraph(h2, style="Heading 2")
        doc.add_paragraph(body, style="Normal")
        doc.add_paragraph(f"表{i} 第{i}章数据汇总")
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "序号"
        table.rows[0].cells[1].text = "内容"
        table.rows[0].cells[2].text = "备注"
        table.rows[1].cells[0].text = str(i)
        table.rows[1].cells[1].text = f"章节{i}示例数据"
        table.rows[1].cells[2].text = "—"

    doc.add_paragraph("六、附则")
    doc.add_paragraph("本文件自印发之日起施行，由办公室负责解释。")
    doc.add_paragraph("附件3 相关制度汇编")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_chinese_path_fixture(path: Path) -> None:
    """中文深层路径 + 书名号文件名。"""
    doc = Document()
    doc.add_paragraph("重点项目可行性研究报告", style="Title")
    doc.add_paragraph("某市发展和改革委员会", style="Subtitle")
    doc.add_paragraph("一、项目概况", style="Heading 1")
    doc.add_paragraph("（一）建设背景", style="Heading 2")
    doc.add_paragraph(
        "项目位于新城片区，总投资约 3.5 亿元，建设期 24 个月。",
        style="Normal",
    )
    doc.add_paragraph("表1 投资估算表")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "费用类别"
    table.rows[0].cells[1].text = "金额(万元)"
    table.rows[0].cells[2].text = "占比"
    table.rows[1].cells[0].text = "建安费"
    table.rows[1].cells[1].text = "28000"
    table.rows[1].cells[2].text = "80%"
    doc.add_paragraph("图1 区位示意图")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_table_raw_fixture(path: Path) -> None:
    """表格预置 Arial 字体，用于验证 table.enabled=false 时不被改写。"""
    doc = Document()
    doc.add_paragraph("表格禁用测试文档")
    table = doc.add_table(rows=2, cols=2)
    for row_idx, label in enumerate(["保留表头", "保留数据"]):
        cell = table.rows[row_idx].cells[0]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(label)
        run.font.name = "Arial"
        run.font.size = Pt(18)
        table.rows[row_idx].cells[1].text = "X"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def custom_config(wt) -> dict[str, Any]:
    cfg = deepcopy(wt.builtin_default_config())
    cfg["page"]["margin_left"] = 2.8
    cfg["page"]["margin_right"] = 2.6
    cfg["page"]["page_number_size"] = "五号"
    cfg["headings"]["indent_left"] = 0.74
    cfg["headings"]["indent_right"] = 0.5
    cfg["headings"]["body_line_spacing"] = 30
    cfg["table"]["border_pt"] = 1.0
    cfg["table"]["row_height_cm"] = 0.75
    cfg["other"]["attachment_enabled"] = True
    return cfg


# ---------------------------------------------------------------------------
# 校验逻辑
# ---------------------------------------------------------------------------


def verify_common(
    scenario: str,
    doc: Document,
    config: dict[str, Any],
    report: TestReport,
    wt,
) -> None:
    section = doc.sections[0]
    page = config["page"]
    headings = config["headings"]
    other = config["other"]

    report.add(
        scenario,
        "页脚 PAGE 字段",
        footer_has_page_field(section),
    )
    fn_font, fn_size = footer_page_run_props(section)
    report.add(
        scenario,
        "页码字体",
        font_matches(fn_font, page["page_number_font"]),
        f"实际 {fn_font}",
    )
    expected_fn_pt = wt.SIZE_MAP[page["page_number_size"]]
    report.add(
        scenario,
        "页码字号",
        fn_size is not None and nearly(fn_size, expected_fn_pt, 0.6),
        f"期望 {expected_fn_pt}pt，实际 {fn_size}pt",
    )
    report.add(
        scenario,
        "左边距",
        nearly(cm_value(section.left_margin) or 0, page["margin_left"], 0.08),
        f"实际 {cm_value(section.left_margin)}cm",
    )


def verify_basic(scenario: str, doc_path: Path, config: dict[str, Any], report: TestReport) -> None:
    wt = load_word_typeset_module()
    doc = Document(str(doc_path))
    headings = config["headings"]
    table_cfg = config["table"]
    other = config["other"]

    verify_common(scenario, doc, config, report, wt)

    title_p = find_paragraph(doc, lambda p: "工作总结报告" in p.text)
    if title_p:
        font, size, _ = first_run_props(title_p)
        report.add(scenario, "题目标题字体", font_matches(font, headings["title_font"]), f"{font}")
        report.add(
            scenario,
            "题目标题字号",
            nearly(pt_value(title_p.runs[0].font.size) or 0, wt.SIZE_MAP[headings["title_size"]], 0.6),
        )
        report.add(scenario, "题目标题居中", title_p.alignment == WD_ALIGN_PARAGRAPH.CENTER)

    body = find_paragraph(doc, lambda p: "English 123" in p.text)
    if body:
        font, size, _ = first_run_props(body)
        report.add(scenario, "正文字体", font_matches(font, headings["body_font"]))
        report.add(
            scenario,
            "正文行距",
            nearly(paragraph_line_spacing_pt(body) or 0, headings["body_line_spacing"], 1.0),
        )
        report.add(
            scenario,
            "正文 ascii 字体",
            font_matches(run_ascii_font(body.runs[0]), other["ascii_font"]),
        )

    subtitle_p = find_paragraph(doc, lambda p: p.text.strip() == "某某单位")
    if subtitle_p:
        font, _, _ = first_run_props(subtitle_p)
        report.add(scenario, "副标题字体", font_matches(font, headings["subtitle_font"]), f"{font}")

    h2 = find_paragraph(doc, lambda p: "主要指标" in p.text)
    if h2:
        font, _, bold = first_run_props(h2)
        report.add(scenario, "二级标题字体", font_matches(font, headings["heading2_font"]))
        report.add(scenario, "二级标题不加粗", bold is not True)

    if doc.tables:
        t = doc.tables[0]
        report.add(
            scenario,
            "表格边框",
            nearly(table_border_size_pt(t) or 0, table_cfg["border_pt"], 0.2),
        )
        if table_cfg.get("row_spacing", 0) > 0:
            cell_p = t.rows[1].cells[0].paragraphs[0]
            ls = paragraph_line_spacing_pt(cell_p)
            report.add(
                scenario,
                "表格行距(磅)",
                ls is not None and nearly(ls, table_cfg["row_spacing"], 1.0),
                f"期望 {table_cfg['row_spacing']}pt，实际 {ls}pt",
            )


def verify_complex(scenario: str, doc_path: Path, config: dict[str, Any], report: TestReport) -> None:
    wt = load_word_typeset_module()
    doc = Document(str(doc_path))
    headings = config["headings"]
    table_cfg = config["table"]
    other = config["other"]

    verify_common(scenario, doc, config, report, wt)

    h_heuristic = find_paragraph(doc, lambda p: p.text.strip() == "二、实施路径")
    if h_heuristic:
        font, size, bold = first_run_props(h_heuristic)
        report.add(scenario, "启发式一级标题字体", font_matches(font, headings["heading1_font"]))
        report.add(scenario, "启发式一级标题加粗", bold is True)
        report.add(scenario, "启发式一级大纲级别", outline_level(h_heuristic) == 0)

    h3 = find_paragraph(doc, lambda p: "任务分解" in p.text)
    if h3:
        font, size, _ = first_run_props(h3)
        report.add(scenario, "三级标题按正文字体", font_matches(font, headings["body_font"]))
        report.add(scenario, "三级标题大纲级别", outline_level(h3) == 2)

    mix = find_paragraph(doc, lambda p: "混排测试" in p.text)
    if mix:
        report.add(scenario, "混排段落全部 run 字体", all_runs_east_asia_font(mix, headings["body_font"]))
        report.add(
            scenario,
            "混排保留加粗",
            any(r.bold for r in mix.runs),
        )
        ascii_ok = all(
            font_matches(run_ascii_font(r), other["ascii_font"])
            for r in mix.runs
            if any(ch.isascii() and ch.isalpha() for ch in r.text)
        )
        report.add(scenario, "混排 ascii 字体", ascii_ok)

    cap_space = find_paragraph(doc, lambda p: p.text.startswith("图 2"))
    if cap_space:
        report.add(scenario, "空格图题识别", cap_space.alignment == WD_ALIGN_PARAGRAPH.CENTER)

    report.add(scenario, "文档含两张表", len(doc.tables) >= 2, f"实际 {len(doc.tables)} 张")

    if len(doc.tables) >= 2:
        t2 = doc.tables[1]
        cell = t2.rows[1].cells[1]
        _, size, _ = first_run_props(cell.paragraphs[0])
        report.add(
            scenario,
            "第二表数字单元格字号",
            nearly(size or 0, wt.SIZE_MAP[table_cfg["font_size"]], 0.6),
            f"{size}pt",
        )
        report.add(
            scenario,
            "第二表边框",
            nearly(table_border_size_pt(t2) or 0, table_cfg["border_pt"], 0.2),
        )
        row_h = cm_value(t2.rows[1].height)
        if row_h is not None and table_cfg.get("row_height_cm"):
            report.add(
                scenario,
                "表格行高",
                nearly(row_h, table_cfg["row_height_cm"], 0.15),
                f"{row_h}cm",
            )


def verify_custom(scenario: str, doc_path: Path, config: dict[str, Any], report: TestReport) -> None:
    wt = load_word_typeset_module()
    doc = Document(str(doc_path))
    headings = config["headings"]
    other = config["other"]

    verify_common(scenario, doc, config, report, wt)

    body = find_paragraph(doc, lambda p: "贯彻落实" in p.text)
    if body:
        report.add(
            scenario,
            "正文左缩进",
            nearly(cm_value(body.paragraph_format.left_indent) or 0, headings["indent_left"], 0.08),
            f"{cm_value(body.paragraph_format.left_indent)}cm",
        )
        report.add(
            scenario,
            "正文右缩进",
            nearly(cm_value(body.paragraph_format.right_indent) or 0, headings["indent_right"], 0.08),
        )
        report.add(
            scenario,
            "自定义正文行距",
            nearly(paragraph_line_spacing_pt(body) or 0, headings["body_line_spacing"], 1.0),
        )

    attach = find_paragraph(doc, lambda p: p.text.startswith("附件1"))
    if attach and other.get("attachment_enabled"):
        font, size, bold = first_run_props(attach)
        report.add(scenario, "附件字体", font_matches(font, other["attachment_font"]))
        report.add(scenario, "附件加粗", bold is True)
        expected_pt = wt.SIZE_MAP[other["attachment_size"]]
        report.add(scenario, "附件字号", nearly(size or 0, expected_pt, 0.6))

    attach2 = find_paragraph(doc, lambda p: "附件 2" in p.text)
    if attach2 and other.get("attachment_enabled"):
        _, _, bold2 = first_run_props(attach2)
        report.add(scenario, "空格附件识别并加粗", bold2 is True)

    cap = find_paragraph(doc, lambda p: p.text.startswith("表1"))
    if cap:
        li = cm_value(cap.paragraph_format.left_indent) or 0
        report.add(scenario, "表题无额外缩进", nearly(li, 0, 0.05), f"{li}cm")

    if doc.tables:
        report.add(
            scenario,
            "加粗边框 1pt",
            nearly(table_border_size_pt(doc.tables[0]) or 0, config["table"]["border_pt"], 0.2),
        )


def verify_heuristic(scenario: str, doc_path: Path, config: dict[str, Any], report: TestReport) -> None:
    wt = load_word_typeset_module()
    doc = Document(str(doc_path))
    headings = config["headings"]

    h1 = find_paragraph(doc, lambda p: p.text.startswith("三、"))
    if h1:
        font, _, bold = first_run_props(h1)
        report.add(scenario, "纯文本一级标题", font_matches(font, headings["heading1_font"]))
        report.add(scenario, "纯文本一级加粗", bold is True)

    h2 = find_paragraph(doc, lambda p: p.text.startswith("（四）"))
    if h2:
        font, _, bold = first_run_props(h2)
        report.add(scenario, "纯文本二级标题", font_matches(font, headings["heading2_font"]))
        report.add(scenario, "纯文本二级不加粗", bold is not True)

    body = find_paragraph(doc, lambda p: "普通正文" in p.text)
    if body:
        font, size, _ = first_run_props(body)
        report.add(scenario, "纯文本正文", font_matches(font, headings["body_font"]))
        report.add(
            scenario,
            "纯文本正文字号",
            nearly(size or 0, wt.SIZE_MAP[headings["body_size"]], 0.6),
        )


def verify_table_disabled(
    scenario: str,
    doc_path: Path,
    _config: dict[str, Any],
    report: TestReport,
) -> None:
    doc = Document(str(doc_path))
    if not doc.tables:
        report.add(scenario, "存在表格", False)
        return
    cell = doc.tables[0].rows[0].cells[0]
    font, size, _ = first_run_props(cell.paragraphs[0])
    report.add(scenario, "禁用后保留 Arial", font_matches(font, "Arial"), f"实际 {font}")
    report.add(scenario, "禁用后保留 18pt", nearly(size or 0, 18, 0.6), f"实际 {size}pt")
    report.add(scenario, "禁用后无边框定义", table_border_size_pt(doc.tables[0]) is None)


def verify_batch(scenario: str, paths: list[Path], report: TestReport) -> None:
    report.add(scenario, "批量输出数量", len(paths) == 2, f"{len(paths)}")
    for p in paths:
        report.add(scenario, f"批量文件存在 {p.name}", p.is_file())


def verify_inplace_backup(scenario: str, work: Path, report: TestReport) -> None:
    bak = work.with_suffix(work.suffix + ".bak")
    report.add(scenario, "原地备份文件", bak.is_file(), str(bak))
    if work.is_file():
        doc = Document(str(work))
        report.add(scenario, "原地文件仍可读", len(doc.paragraphs) > 0)


def verify_long_document(
    scenario: str,
    doc_path: Path,
    config: dict[str, Any],
    report: TestReport,
) -> None:
    wt = load_word_typeset_module()
    doc = Document(str(doc_path))
    headings = config["headings"]

    report.add(scenario, "长文五张表", len(doc.tables) == 5, f"实际 {len(doc.tables)}")
    report.add(scenario, "长文章节数", len(doc.paragraphs) >= 20, f"实际 {len(doc.paragraphs)} 段")

    h6 = find_paragraph(doc, lambda p: p.text.strip() == "六、附则")
    if h6:
        font, _, bold = first_run_props(h6)
        report.add(scenario, "第六章启发式一级", font_matches(font, headings["heading1_font"]))
        report.add(scenario, "第六章加粗", bold is True)

    attach = find_paragraph(doc, lambda p: p.text.startswith("附件3"))
    if config["other"].get("attachment_enabled"):
        _, _, bold = first_run_props(attach) if attach else (None, None, None)
        report.add(scenario, "长文附件加粗", attach is not None and bold is True)
    else:
        body = attach and first_run_props(attach)[0]
        report.add(
            scenario,
            "长文附件按正文",
            attach is None or font_matches(body, headings["body_font"]),
        )

    last_table = doc.tables[-1] if doc.tables else None
    if last_table:
        cell = last_table.rows[1].cells[1]
        _, size, _ = first_run_props(cell.paragraphs[0])
        report.add(
            scenario,
            "末表字号",
            nearly(size or 0, wt.SIZE_MAP[config["table"]["font_size"]], 0.6),
        )


def verify_chinese_path(
    scenario: str,
    doc_path: Path,
    config: dict[str, Any],
    report: TestReport,
) -> None:
    wt = load_word_typeset_module()
    doc = Document(str(doc_path))

    report.add(scenario, "中文路径输出可读", doc_path.is_file())
    report.add(scenario, "输出文件名含书名号", "《" in doc_path.name)

    title = find_paragraph(doc, lambda p: "可行性研究报告" in p.text)
    if title:
        font, size, _ = first_run_props(title)
        report.add(
            scenario,
            "报告标题字体",
            font_matches(font, config["headings"]["title_font"]),
        )
        report.add(
            scenario,
            "报告标题字号",
            nearly(size or 0, wt.SIZE_MAP[config["headings"]["title_size"]], 0.6),
        )

    body = find_paragraph(doc, lambda p: "新城片区" in p.text)
    if body:
        report.add(
            scenario,
            "正文行距",
            nearly(
                paragraph_line_spacing_pt(body) or 0,
                config["headings"]["body_line_spacing"],
                1.0,
            ),
        )

    if doc.tables:
        report.add(
            scenario,
            "投资表边框",
            nearly(
                table_border_size_pt(doc.tables[0]) or 0,
                config["table"]["border_pt"],
                0.2,
            ),
        )


def cleanup_legacy_fixtures() -> None:
    for base in (FIXTURES_DIR, ARTIFACTS_DIR):
        for name in LEGACY_FILES:
            p = base / name
            if p.exists():
                p.unlink()
        legacy_batch = base / "batch-run"
        if legacy_batch.is_dir():
            shutil.rmtree(legacy_batch, ignore_errors=True)


def verify_text_mode(scenario: str, out_dir: Path, config: dict[str, Any], report: TestReport) -> None:
    out = out_dir / F["文本_出"]
    payload = {
        "mode": "text",
        "text": (
            "一、测试章节\n"
            "\n\n"
            "正文含 English 99 与表意文字。\n"
            "## Markdown 二级\n"
            "表4 文本模式表题\n"
            "附件1 测试附件"
        ),
        "output_path": str(out),
        "config": config,
    }
    try:
        run_typeset_payload(payload)
        report.add(scenario, "文本模式成功", out.is_file())
        if out.is_file():
            doc = Document(str(out))
            h1 = find_paragraph(doc, lambda p: p.text.startswith("一、"))
            report.add(scenario, "文本模式识别一级标题", h1 is not None)
            md_h2 = find_paragraph(doc, lambda p: "Markdown 二级" in p.text)
            if md_h2:
                font, _, bold = first_run_props(md_h2)
                report.add(
                    scenario,
                    "MD二级标题字体",
                    font_matches(font, config["headings"]["heading2_font"]),
                )
                report.add(scenario, "MD二级标题不加粗", bold is not True)
            cap = find_paragraph(doc, lambda p: p.text.startswith("表4"))
            report.add(
                scenario,
                "文本模式表题居中",
                cap is not None and cap.alignment == WD_ALIGN_PARAGRAPH.CENTER,
            )
            attach = find_paragraph(doc, lambda p: p.text.startswith("附件1"))
            if config["other"].get("attachment_enabled"):
                report.add(
                    scenario,
                    "文本模式附件格式化",
                    attach is not None and first_run_props(attach)[2] is True,
                )
            # 合并空行：连续两个空行应只保留一个空段
            empty_count = sum(1 for p in doc.paragraphs if not p.text.strip())
            report.add(
                scenario,
                "TXT合并连续空行",
                empty_count <= 2,
                f"空段落数 {empty_count}",
            )
    except Exception as exc:
        report.add(scenario, "文本模式成功", False, str(exc))


# ---------------------------------------------------------------------------
# 场景执行
# ---------------------------------------------------------------------------


def run_scenario_basic(report: TestReport, wt) -> None:
    scenario = "基础"
    inp = ARTIFACTS_DIR / F["基础_前"]
    out = ARTIFACTS_DIR / F["基础_后"]
    build_basic_fixture(inp)
    if out.exists():
        out.unlink()
    run_typeset_file(inp, out, wt.builtin_default_config())
    report.add(scenario, "排版成功", out.is_file())
    verify_basic(scenario, out, wt.builtin_default_config(), report)


def run_scenario_complex(report: TestReport, wt) -> None:
    scenario = "复杂"
    inp = ARTIFACTS_DIR / F["复杂_前"]
    out = ARTIFACTS_DIR / F["复杂_后"]
    build_complex_fixture(inp)
    if out.exists():
        out.unlink()
    cfg = wt.builtin_default_config()
    run_typeset_file(inp, out, cfg)
    report.add(scenario, "排版成功", out.is_file())
    verify_complex(scenario, out, cfg, report)


def run_scenario_custom(report: TestReport, wt) -> None:
    scenario = "自定义配置"
    inp = ARTIFACTS_DIR / F["复杂_前"]
    out = ARTIFACTS_DIR / F["自定义_后"]
    if not inp.is_file():
        build_complex_fixture(inp)
    cfg = custom_config(wt)
    if out.exists():
        out.unlink()
    run_typeset_file(inp, out, cfg)
    report.add(scenario, "排版成功", out.is_file())
    verify_custom(scenario, out, cfg, report)


def run_scenario_heuristic(report: TestReport, wt) -> None:
    scenario = "启发式"
    inp = ARTIFACTS_DIR / F["启发_前"]
    out = ARTIFACTS_DIR / F["启发_后"]
    build_heuristic_fixture(inp)
    cfg = wt.builtin_default_config()
    if out.exists():
        out.unlink()
    run_typeset_file(inp, out, cfg)
    report.add(scenario, "排版成功", out.is_file())
    verify_heuristic(scenario, out, cfg, report)


def run_scenario_batch(report: TestReport, wt) -> None:
    scenario = "批量"
    batch_dir = ARTIFACTS_DIR / DIR_BATCH
    batch_dir.mkdir(parents=True, exist_ok=True)
    a_in = batch_dir / F["批量_A"]
    b_in = batch_dir / F["批量_B"]
    build_basic_fixture(a_in)
    build_heuristic_fixture(b_in)
    result = run_typeset_payload(
        {
            "mode": "batch",
            "input_paths": [str(a_in), str(b_in)],
            "in_place": True,
            "config": wt.builtin_default_config(),
        }
    )
    report.add(scenario, "批量脚本返回2条", len(result.get("results", [])) == 2)
    verify_batch(scenario, [a_in, b_in], report)
    doc_a = Document(str(a_in))
    report.add(
        scenario,
        "批量文件A内容完整",
        find_paragraph(doc_a, lambda p: "工作总结报告" in p.text) is not None,
    )
    doc_b = Document(str(b_in))
    report.add(
        scenario,
        "批量文件B标题已排版",
        find_paragraph(doc_b, lambda p: p.text.startswith("三、")) is not None,
    )


def run_scenario_table_disabled(report: TestReport, wt) -> None:
    scenario = "表格禁用"
    inp = ARTIFACTS_DIR / F["表格禁_原"]
    out = ARTIFACTS_DIR / F["表格禁_后"]
    build_table_raw_fixture(inp)
    cfg = wt.builtin_default_config()
    cfg["table"]["enabled"] = False
    if out.exists():
        out.unlink()
    run_typeset_file(inp, out, cfg)
    report.add(scenario, "排版成功", out.is_file())
    verify_table_disabled(scenario, out, cfg, report)


def run_scenario_long_document(report: TestReport, wt) -> None:
    scenario = "公文长文"
    inp = ARTIFACTS_DIR / F["长文_前"]
    out = ARTIFACTS_DIR / F["长文_后"]
    build_long_document_fixture(inp)
    cfg = custom_config(wt)
    if out.exists():
        out.unlink()
    run_typeset_file(inp, out, cfg)
    report.add(scenario, "排版成功", out.is_file())
    verify_long_document(scenario, out, cfg, report)


def run_scenario_chinese_path(report: TestReport, wt) -> None:
    scenario = "中文路径"
    cn_dir = ARTIFACTS_DIR / DIR_CN_PATH
    inp = cn_dir / F["中文路径_文件"]
    out = cn_dir / F["中文路径_后"]
    build_chinese_path_fixture(inp)
    cfg = wt.builtin_default_config()
    if out.exists():
        out.unlink()
    run_typeset_file(inp, out, cfg)
    report.add(scenario, "排版成功", out.is_file())
    report.add(scenario, "输入含书名号", "《" in inp.name and "》" in inp.name)
    report.add(scenario, "路径含中文目录", "某市发改委" in str(inp))
    verify_chinese_path(scenario, out, cfg, report)


def run_scenario_inplace(report: TestReport, wt) -> None:
    scenario = "原地备份"
    work = ARTIFACTS_DIR / F["原地_稿"]
    bak = work.with_suffix(work.suffix + ".bak")
    build_basic_fixture(work)
    if bak.exists():
        bak.unlink()
    run_typeset_file(work, work, wt.builtin_default_config(), in_place=True)
    verify_inplace_backup(scenario, work, report)
    if bak.exists():
        bak.unlink()


def main() -> int:
    wt = load_word_typeset_module()
    report = TestReport()

    print("==> Word 批量排版自动化测试（多场景 · 中文文件名）")
    ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)
    cleanup_legacy_fixtures()

    scenarios: list[tuple[str, Callable[[], None]]] = [
        ("基础样例", lambda: run_scenario_basic(report, wt)),
        ("复杂样例", lambda: run_scenario_complex(report, wt)),
        ("自定义配置", lambda: run_scenario_custom(report, wt)),
        ("启发式标题", lambda: run_scenario_heuristic(report, wt)),
        ("公文长文", lambda: run_scenario_long_document(report, wt)),
        ("中文路径", lambda: run_scenario_chinese_path(report, wt)),
        ("批量排版", lambda: run_scenario_batch(report, wt)),
        ("表格禁用", lambda: run_scenario_table_disabled(report, wt)),
        ("原地备份", lambda: run_scenario_inplace(report, wt)),
    ]

    text_dir = ARTIFACTS_DIR / "文本模式"
    text_dir.mkdir(parents=True, exist_ok=True)
    verify_text_mode("文本模式", text_dir, wt.builtin_default_config(), report)

    for label, fn in scenarios:
        print(f"==> 场景: {label}")
        try:
            fn()
        except Exception as exc:
            report.add(label, "场景执行", False, str(exc))
            print(f"    场景失败: {exc}")

    summary = report.dump()
    failed = summary["failed"]

    print("")
    print(f"检查结果: {summary['total'] - len(failed)}/{summary['total']} 通过")
    print("分场景:")
    for name, info in summary["scenarios"].items():
        mark = "✓" if not info["failed"] else "✗"
        print(f"  {mark} {name}: {info['passed']}/{info['total']}")

    for c in summary["checks"]:
        if not c["ok"]:
            detail = f" — {c['detail']}" if c["detail"] else ""
            print(f"  ✗ [{c['scenario']}] {c['name']}{detail}")

    result_path = ARTIFACTS_DIR / "last-test-report.json"
    result_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n报告: {result_path}")

    if summary["passed"]:
        print("\n全部通过。")
        return 0
    print(f"\n失败 {len(failed)} 项")
    return 1


if __name__ == "__main__":
    sys.exit(main())
