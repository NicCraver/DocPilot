#!/usr/bin/env python3
"""从用户提供的《选煤固废资源化利用研究进展》PDF 生成复杂期刊 docx 并排版测试。"""
from __future__ import annotations

import importlib.util
import json
import re
import subprocess
import sys
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent.parent
PDF_PATH = Path("/Users/nic/Downloads/选煤固废资源化利用研究进展_1.pdf")
SCRIPT = ROOT / "scripts" / "word-typeset.py"
DATA = ROOT / "scripts" / "word-typeset-test-data"
ARTIFACTS = ROOT / "scripts" / "test-artifacts" / "word-typeset"
ASSETS = DATA / "assets" / "coal-paper-full"
DOC_IN = ARTIFACTS / "选煤固废-复杂论文-排版前.docx"
DOC_OUT = ARTIFACTS / "选煤固废-复杂论文-排版后.docx"

HEADER_NOISE = re.compile(
    r"^(doi:|第\d+卷|Vol\.|COAL|ENGINEE|煤$|炭$|工$|程$|专家论坛|\d{1,3}$|"
    r"2024\s*年|收稿日期|作者简介|引用格式|Research progress|DONG Xianshu|"
    r"College of Mining|Abstract:|Keywords:|摘$|要$)",
    re.I,
)
SECTION_H1 = re.compile(r"^(\d{1,2})\s+([\u4e00-\u9fff].+)")
SECTION_H2 = re.compile(r"^(\d+)\.\s*(\d+)\s+([\u4e00-\u9fff].+)")
TABLE_CN = re.compile(r"^表\s*(\d+)\s*(.*)")
TABLE_EN = re.compile(r"^Table\s+(\d+)\s*(.*)", re.I)
FIGURE_CN = re.compile(r"^图\s*(\d+)\s*(.*)")
CITATION_INLINE = re.compile(r"［(\d+(?:[，,\-\s]\d+)*)］|\[(\d+(?:[，,\-\s]\d+)*)\]")

H1_TITLES = (
    "选煤固废的物理化学性质",
    "煤矸石资源化利用研究进展",
    "尾煤资源化利用研究进展",
    "选煤固废资源化利用面临的问题",
    "结语及展望",
)
H2_TITLES = frozenset(
    {
        "煤矸石",
        "尾煤",
        "煤矸石综合利用",
        "尾煤吸附材料",
        "尾煤充填材料",
        "尾煤土壤改良剂",
    }
)
TABLE_ONLY = re.compile(r"^表\s*(\d+)\s*$")
FIGURE_ONLY = re.compile(r"^图\s*(\d+)\s*(.*)$")
REF_START = re.compile(r"^[\[［]1[\]］]")


def load_module():
    spec = importlib.util.spec_from_file_location("word_typeset", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


def find_cjk_font(size: int = 16):
    for path in (
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "C:/Windows/Fonts/msyh.ttc",
    ):
        if Path(path).is_file():
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                continue
    return ImageFont.load_default()


def draw_figure(path: Path, title: str, kind: str, idx: int) -> None:
    w, h = 680, 320
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = find_cjk_font(18)
    draw.text((24, 16), title, fill=(30, 64, 175), font=title_font)
    colors = [(59, 130, 246), (34, 197, 94), (245, 158, 11), (168, 85, 247)]
    c = colors[idx % len(colors)]
    fill = tuple(min(255, int(x * 0.35 + 180)) for x in c)
    if kind == "flow":
        boxes = ["原料", "处理", "产品", "应用"]
        for i, label in enumerate(boxes):
            x0 = 40 + i * 155
            draw.rounded_rectangle((x0, 100, x0 + 130, 180), radius=8, fill=fill, outline=c, width=2)
            bbox = draw.textbbox((0, 0), label, font=find_cjk_font(14))
            tw = bbox[2] - bbox[0]
            draw.text((x0 + (130 - tw) / 2, 130), label, fill=(33, 37, 41), font=find_cjk_font(14))
    elif kind == "chart":
        for i, hval in enumerate([60, 85, 45, 70, 55]):
            x0 = 80 + i * 100
            bh = int(140 * hval / 100)
            draw.rectangle((x0, 260 - bh, x0 + 60, 260), fill=c)
    else:
        draw.ellipse((240, 90, 440, 240), outline=c, width=3)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, format="PNG")


def normalize_citations(text: str) -> str:
    return CITATION_INLINE.sub(lambda m: f"[{m.group(1) or m.group(2)}]", text)


def split_h1_line(text: str) -> tuple[str, str]:
    for title in H1_TITLES:
        if text == title:
            return title, ""
        if text.startswith(title):
            return title, text[len(title) :].lstrip()
    if text and text[0].isdigit() and " " in text:
        num, rest = text.split(" ", 1)
        if num.isdigit() and rest[:1].isalpha() or rest[:1] == "\u4e00":
            return text, ""
    return text, ""


class HeadingNumberer:
    def __init__(self) -> None:
        self.h1 = 0
        self.h2 = 0

    def h1_text(self, title: str) -> str:
        title = re.sub(r"^\d+\s+", "", title.strip())
        self.h1 += 1
        self.h2 = 0
        return f"{self.h1} {title}"

    def h2_text(self, title: str) -> str:
        title = re.sub(r"^\d+\.\d+\s+", "", title.strip())
        self.h2 += 1
        return f"{self.h1}.{self.h2} {title}"


def extract_page_lines(page: fitz.Page) -> list[dict[str, float | str | int]]:
    mid = page.rect.width / 2
    items: list[dict[str, float | str | int]] = []
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:
            continue
        for line in block["lines"]:
            spans = line["spans"]
            text = "".join(s["text"] for s in spans).strip()
            if not text:
                continue
            x0 = min(s["bbox"][0] for s in spans)
            x1 = max(s["bbox"][2] for s in spans)
            y0 = min(s["bbox"][1] for s in spans)
            max_size = max(s["size"] for s in spans)
            col = 0 if (x0 + x1) / 2 < mid else 1
            items.append({"text": text, "y": y0, "col": col, "size": max_size})
    items.sort(key=lambda i: (int(i["col"]), float(i["y"])))
    return items


def is_table_data_line(text: str, size: float) -> bool:
    if size > 9.2:
        return False
    if re.match(r"^[\d.\s～\-~%mg/kgμ]+$", text, re.I):
        return True
    if re.match(r"^[A-Za-z]{1,3}(\s|$)", text) and len(text) < 20:
        return True
    return False


def classify_line(text: str, size: float) -> str:
    if not text or HEADER_NOISE.match(text):
        return "skip"
    if REF_START.match(text):
        return "refs"
    if size >= 11.5:
        return "h1"
    for title in H1_TITLES:
        if text.startswith(title):
            return "h1"
    if text in H2_TITLES:
        return "h2"
    if TABLE_ONLY.match(text) or (TABLE_CN.match(text) and size <= 10):
        return "table"
    if FIGURE_ONLY.match(text) or (FIGURE_CN.match(text) and size <= 10):
        return "figure"
    if re.match(r"^Table\s+\d+", text, re.I):
        return "table_en"
    if re.match(r"^Fig\.?\s+\d+", text, re.I):
        return "figure_en"
    if is_table_data_line(text, size):
        return "skip"
    if size <= 8.6 and len(text) < 40:
        return "caption_bits"
    return "body"


def extract_structured_body(pdf: fitz.Document) -> list[dict[str, str]]:
    raw: list[dict[str, float | str | int]] = []
    for pno in range(1, len(pdf)):
        raw.extend(extract_page_lines(pdf[pno]))

    output: list[dict[str, str]] = []
    buf = ""
    i = 0
    refs = False

    def flush_body() -> None:
        nonlocal buf
        if buf.strip():
            output.append({"type": "body", "text": normalize_citations(buf.strip())})
        buf = ""

    while i < len(raw):
        item = raw[i]
        text = str(item["text"]).strip()
        size = float(item["size"])
        if refs:
            i += 1
            continue

        kind = classify_line(text, size)
        if kind == "refs":
            flush_body()
            refs = True
            i += 1
            continue
        if kind == "skip" or kind == "caption_bits":
            i += 1
            continue

        if kind == "h1":
            flush_body()
            title, rest = split_h1_line(text)
            output.append({"type": "h1", "text": title})
            if rest:
                buf = rest
            i += 1
            continue

        if kind == "h2":
            flush_body()
            output.append({"type": "h2", "text": text})
            i += 1
            continue

        if kind == "table":
            flush_body()
            m = TABLE_CN.match(text) or TABLE_ONLY.match(text)
            tid = m.group(1) if m else re.search(r"\d+", text).group()  # type: ignore[union-attr]
            cn_title = (m.group(2).strip() if m and m.lastindex and m.lastindex >= 2 else "") or ""
            j = i + 1
            if not cn_title:
                parts: list[str] = []
                while j < len(raw):
                    nxt = str(raw[j]["text"]).strip()
                    nsize = float(raw[j]["size"])
                    if re.match(r"^Table\s+\d+", nxt, re.I):
                        break
                    if classify_line(nxt, nsize) not in ("caption_bits", "skip") and nsize > 9:
                        break
                    if nxt and not re.match(r"^表\s*\d", nxt) and nsize <= 9.5:
                        parts.append(nxt)
                    j += 1
                cn_title = " ".join(parts).strip()
            en_title = ""
            if j < len(raw) and re.match(r"^Table\s+\d+", str(raw[j]["text"]).strip(), re.I):
                en_title = re.sub(r"^Table\s+\d+\s*", "", str(raw[j]["text"]).strip(), flags=re.I)
                j += 1
            while j < len(raw) and is_table_data_line(str(raw[j]["text"]).strip(), float(raw[j]["size"])):
                j += 1
            output.append({"type": "table", "id": tid, "cn": cn_title, "en": en_title})
            i = j
            continue

        if kind == "figure":
            flush_body()
            m = FIGURE_CN.match(text) or FIGURE_ONLY.match(text)
            fid = m.group(1) if m else re.search(r"\d+", text).group()  # type: ignore[union-attr]
            cn_title = (m.group(2).strip() if m and m.lastindex and m.lastindex >= 2 else "") or ""
            j = i + 1
            if not cn_title:
                parts = []
                while j < len(raw):
                    nxt = str(raw[j]["text"]).strip()
                    nsize = float(raw[j]["size"])
                    if re.match(r"^Fig\.?\s+\d+", nxt, re.I):
                        break
                    if classify_line(nxt, nsize) in ("h1", "h2", "body", "table", "figure") and nsize > 9:
                        break
                    if nxt and not re.match(r"^图\s*\d", nxt) and nsize <= 9.8:
                        parts.append(nxt)
                    j += 1
                cn_title = " ".join(parts).strip()
            en_title = ""
            if j < len(raw) and re.match(r"^Fig\.?\s+\d+", str(raw[j]["text"]).strip(), re.I):
                en_title = re.sub(r"^Fig\.?\s+\d+[A-Za-z]?[\s\-－]*", "", str(raw[j]["text"]).strip(), flags=re.I)
                j += 1
            output.append({"type": "figure", "id": fid, "cn": cn_title, "en": en_title})
            i = j
            continue

        if kind in ("table_en", "figure_en"):
            i += 1
            continue

        if buf:
            if buf.endswith("-") or re.search(r"[a-zA-Z]$", buf):
                buf = buf.rstrip("-") + text
            else:
                buf += text
        else:
            buf = text
        if buf.endswith(("。", "！", "？", "；", ":", "：")) or len(buf) > 220:
            flush_body()
        i += 1

    flush_body()
    return output


def extract_body_from_pdf(pdf: fitz.Document) -> list[str]:
    """兼容旧接口：返回扁平文本行（含编号标题）。"""
    numberer = HeadingNumberer()
    lines: list[str] = []
    for item in extract_structured_body(pdf):
        t = item["type"]
        if t == "body":
            lines.append(item["text"])
        elif t == "h1":
            lines.append(numberer.h1_text(item["text"]))
        elif t == "h2":
            lines.append(numberer.h2_text(item["text"]))
        elif t == "table":
            cap = f"表{item['id']} {item['cn']}".strip()
            lines.append(cap)
            if item.get("en"):
                lines.append(f"Table {item['id']} {item['en']}".strip())
        elif t == "figure":
            cap = f"图{item['id']} {item['cn']}".strip()
            lines.append(cap)
    return lines


def extract_front_matter(pdf: fitz.Document) -> dict[str, str]:
    page0 = pdf[0]
    text = page0.get_text()
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    def find_after(prefix: str) -> str:
        for ln in lines:
            if ln.startswith(prefix):
                return ln
        return ""

    abstract_parts: list[str] = []
    in_abs = False
    for ln in lines:
        if ln == "摘" or ln.startswith("摘"):
            in_abs = True
            if "要" in ln and "：" in ln:
                abstract_parts.append(ln.split("：", 1)[-1])
            continue
        if in_abs:
            if ln.startswith("关键词"):
                break
            abstract_parts.append(ln)
    abstract = normalize_citations("".join(abstract_parts))
    if abstract.startswith("要"):
        abstract = abstract.lstrip("要:").lstrip("要：").strip()

    doi = find_after("doi") or "doi：10.11799/ce202410012"
    doi = re.sub(r"\s+", "", doi.replace("doi:", "doi："))

    return {
        "doi": doi,
        "title": "选煤固废资源化利用研究进展",
        "author": "董宪姝",
        "affiliation": "（太原理工大学 矿业工程学院，山西 太原 030024）",
        "abstract": f"摘要：{abstract}" if not abstract.startswith("摘要") else abstract,
        "keywords": find_after("关键词") or "关键词：煤炭洗选；固体废弃物；物理化学性质；综合利用",
        "metadata": "中图分类号：TD94    文献标识码：A    文章编号：1671-0959（2024）10-0108-14",
        "abstract_en": (
            "Abstract: Coal beneficiation is the premise and foundation for clean and efficient "
            "utilization of coal, and the solid waste generated from the beneficiation process "
            "has been increasing year by year."
        ),
        "keywords_en": "Keywords: coal beneficiation; solid waste; physico-chemical properties; comprehensive utilization",
        "received": "收稿日期：2024-06-15",
        "bio": "作者简介：董宪姝（1964—），女，辽宁葫芦岛人，教授，博士生导师，主要研究方向为煤炭资源清洁高效洗选与煤基固废梯级利用等。",
        "cite": "引用格式：董宪姝. 选煤固废资源化利用研究进展[J]. 煤炭工程，2024，56（10）：108-121.",
    }


TABLE_SAMPLES: dict[str, tuple[list[str], list[list[str]]]] = {
    "1": (
        ["来源", "表观密度/(kg·m⁻³)", "堆积密度/(kg·m⁻³)", "吸水率/%"],
        [
            ["大同", "2100～2400", "1300～1500", "1.7～3.9"],
            ["阳泉", "2200～2500", "1400～1600", "2.1～4.5"],
            ["淮南", "2000～2300", "1250～1450", "1.8～5.2"],
        ],
    ),
    "2": (
        ["元素", "Al", "Fe", "S", "Mg", "Cu", "Li"],
        [
            ["含量/(mg·kg⁻¹)", "65000", "42000", "8500", "12000", "120", "85"],
            ["富集特征", "高", "中", "中", "中", "低", "低"],
        ],
    ),
    "3": (
        ["项目", "灰分/%", "挥发分/%", "固定碳/%", "发热量/(MJ·kg⁻¹)"],
        [
            ["尾煤 A", "45.2", "18.6", "36.2", "12.5"],
            ["尾煤 B", "52.8", "15.3", "31.9", "10.8"],
        ],
    ),
}


def add_table(doc: Document, table_id: str, cn_title: str, en_title: str | None) -> None:
    doc.add_paragraph(f"表{table_id} {cn_title}".strip())
    if en_title:
        doc.add_paragraph(f"Table {table_id} {en_title}".strip())
    hdr, rows = TABLE_SAMPLES.get(table_id, (["列1", "列2", "列3"], [["—", "—", "—"]]))
    table = doc.add_table(rows=1 + len(rows), cols=len(hdr))
    for c, h in enumerate(hdr):
        table.rows[0].cells[c].text = h
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val


def add_figure(doc: Document, fig_id: str, caption: str, png: Path) -> None:
    if png.is_file():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(png), width=Inches(4.5))
    doc.add_paragraph(f"图{fig_id} {caption}".strip())


def build_complex_docx(pdf_path: Path, out_path: Path) -> tuple[int, int, int]:
    if not pdf_path.is_file():
        raise FileNotFoundError(f"未找到 PDF: {pdf_path}")

    pdf = fitz.open(str(pdf_path))
    fm = extract_front_matter(pdf)
    structured = extract_structured_body(pdf)
    pdf.close()

    ASSETS.mkdir(parents=True, exist_ok=True)
    fig_count = 0
    tbl_count = 0
    numberer = HeadingNumberer()

    doc = Document()
    doc.add_paragraph(fm["doi"])
    doc.add_paragraph(fm["title"], style="Title")
    doc.add_paragraph(fm["author"], style="Subtitle")
    doc.add_paragraph(fm["affiliation"])
    doc.add_paragraph(normalize_citations(fm["abstract"]))
    doc.add_paragraph(fm["keywords"])
    doc.add_paragraph(fm["metadata"])
    doc.add_paragraph(fm["abstract_en"])
    doc.add_paragraph(fm["keywords_en"])
    doc.add_paragraph(fm["received"])
    doc.add_paragraph(fm["bio"])
    doc.add_paragraph(fm["cite"])

    for item in structured:
        kind = item["type"]
        if kind == "h1":
            doc.add_paragraph(numberer.h1_text(item["text"]))
            continue
        if kind == "h2":
            doc.add_paragraph(numberer.h2_text(item["text"]))
            continue
        if kind == "table":
            cn = item["cn"] or f"数据表{item['id']}"
            en = item.get("en") or None
            add_table(doc, item["id"], cn, en)
            tbl_count += 1
            continue
        if kind == "figure":
            fig_id = item["id"]
            caption = item["cn"] or item.get("en") or f"示意图{fig_id}"
            kind_img = "flow" if int(fig_id) % 3 == 1 else "chart" if int(fig_id) % 3 == 2 else "schematic"
            png = ASSETS / f"fig{fig_id}.png"
            draw_figure(png, caption, kind_img, int(fig_id))
            add_figure(doc, fig_id, caption, png)
            fig_count += 1
            continue
        if kind == "body":
            doc.add_paragraph(item["text"])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    char_count = sum(len(re.sub(r"\s", "", p.text)) for p in doc.paragraphs)
    return char_count, tbl_count, fig_count


def run_typeset(inp: Path, out: Path, config: dict) -> None:
    py = ROOT / ".venv" / "bin" / "python3"
    if not py.is_file():
        py = Path(sys.executable)
    payload = {
        "mode": "batch",
        "input_paths": [str(inp)],
        "output_path": str(out),
        "in_place": False,
        "config": config,
    }
    proc = subprocess.run(
        [str(py), str(SCRIPT), json.dumps(payload, ensure_ascii=False)],
        capture_output=True,
        text=True,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr or proc.stdout)
    result = json.loads(proc.stdout.strip())
    if not result.get("ok"):
        raise RuntimeError(result.get("error", "排版失败"))


def verify(out: Path) -> list[tuple[str, bool, str]]:
    doc = Document(str(out))
    checks: list[tuple[str, bool, str]] = []
    cols = []
    for s in doc.sections:
        c = s._sectPr.find(qn("w:cols"))
        cols.append(int(c.get(qn("w:num"))) if c is not None and c.get(qn("w:num")) else 1)
    checks.append(("双栏分节", len(cols) >= 2 and cols[-1] == 2, str(cols)))

    chars = sum(len(re.sub(r"\s", "", p.text)) for p in doc.paragraphs)
    checks.append(("论文字数", chars >= 8000, f"{chars} 字"))

    checks.append(("表格", len(doc.tables) >= 2, f"{len(doc.tables)} 张"))
    inline = sum(len(p._element.xpath(".//wp:inline")) for p in doc.paragraphs)
    checks.append(("插图", inline >= 3, f"{inline} 张"))

    hdr = " ".join(p.text for s in doc.sections for p in s.header.paragraphs)
    checks.append(("页眉", "煤炭工程" in hdr.replace(" ", ""), hdr[:60]))

    cite = next((p for p in doc.paragraphs if re.search(r"\[\d", p.text)), None)
    sup = cite and any(r.font.superscript for r in cite.runs if r.text and "[" in r.text)
    checks.append(("引文上标", bool(sup), "ok" if sup else "无"))

    h1 = sum(
        1
        for p in doc.paragraphs
        if re.match(r"^[1-9]\d?\s+[\u4e00-\u9fff]{2}", (p.text or "").strip())
    )
    h2 = sum(
        1
        for p in doc.paragraphs
        if re.match(r"^\d+\.\d+\s+[\u4e00-\u9fff]", (p.text or "").strip())
    )
    checks.append(("章节标题", h1 >= 3 and h2 >= 5, f"h1={h1} h2={h2}"))

    return checks


def main() -> int:
    wt = load_module()
    print(f"==> 从 PDF 生成复杂期刊论文 docx")
    print(f"    PDF: {PDF_PATH}")

    chars, tables, figs = build_complex_docx(PDF_PATH, DOC_IN)
    print(f"    生成: {DOC_IN}")
    print(f"    规模: 约 {chars} 字，表格 {tables} 张，配图 {figs} 张")

    if DOC_OUT.exists():
        DOC_OUT.unlink()
    config = wt.journal_default_config()
    run_typeset(DOC_IN, DOC_OUT, config)
    print(f"    排版: {DOC_OUT}")

    doc = Document(str(DOC_OUT))
    print(f"    排版后: {len(doc.paragraphs)} 段，{len(doc.tables)} 表，{len(doc.sections)} 节")

    failed = 0
    print("\n==> 校验")
    for name, ok, detail in verify(DOC_OUT):
        print(f"    {'✓' if ok else '✗'} {name}: {detail}")
        if not ok:
            failed += 1
    if failed:
        return 1
    print("\n完成。请用 Word 打开「选煤固废-复杂论文-排版后.docx」查看。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
