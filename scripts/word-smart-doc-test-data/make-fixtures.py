#!/usr/bin/env python3
"""用 python-docx 生成 word-smart-doc 测试样例：两份结构不同的模板 + 一份内容。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

def _set_normal_body(doc: Document, font_ea: str, size_pt: float) -> None:
    """在 Normal 样式上定义正文宋体小四，占位段继承样式、不显式设 run 字号。"""
    normal = doc.styles["Normal"]
    normal.font.name = font_ea
    normal.font.size = Pt(size_pt)
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font_ea)

HERE = Path(__file__).resolve().parent


def _heading(doc, text: str, size: int, font_ea: str, bold: bool):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_ea
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        font_ea,
    )
    return p


def make_year_end_template():
    doc = Document()
    _set_normal_body(doc, "宋体", 12)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("2025 年终总结报告")
    r.bold = True
    r.font.size = Pt(22)

    guide = doc.add_paragraph()
    gr = guide.add_run("（请在此处填写内容，红色说明使用后删除）")
    gr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    gr.italic = True

    for h in ["前言", "工作成果", "工作复盘", "未来展望", "结语"]:
        _heading(doc, h, 16, "黑体", True)
        doc.add_paragraph("（此处填写正文）")

    doc.add_paragraph("汇报人：")
    doc.add_paragraph("日期：")
    out = HERE / "year-end-template.docx"
    doc.save(str(out))
    return out


def make_realistic_body_style_template():
    """模拟常见企业模板：Normal=四号 14pt，占位段挂「正文」样式=小四 12pt。"""
    doc = Document()
    _set_normal_body(doc, "宋体", 14)
    body_style = doc.styles.add_style("正文", WD_STYLE_TYPE.PARAGRAPH)
    body_style.base_style = doc.styles["Normal"]
    body_style.font.name = "宋体"
    body_style.font.size = Pt(12)
    rpr = body_style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")

    for h in ["前言", "工作成果", "结语"]:
        _heading(doc, h, 16, "黑体", True)
        doc.add_paragraph("（此处填写正文）", style="正文")

    out = HERE / "realistic-body-style.docx"
    doc.save(str(out))
    return out


def make_gov_template():
    doc = Document()
    _set_normal_body(doc, "仿宋_GB2312", 16)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("关于加强本市数据治理工作的通知")
    r.font.size = Pt(22)
    r.font.name = "方正小标宋简体"

    for h in ["一、总体要求", "二、主要任务", "三、保障措施"]:
        _heading(doc, h, 16, "黑体", False)
        doc.add_paragraph("（此处填写正文）")
    out = HERE / "gov-template.docx"
    doc.save(str(out))
    return out


def make_content_md():
    md = """# 前言
今年是公司高速发展的一年，团队完成了多项关键目标。

# 工作成果
完成核心系统重构，性能提升 40%。
上线智能模板库功能，获得客户好评。

# 工作复盘
项目排期偏紧，需要加强前期评估。

# 未来展望
2026 年将聚焦智能化与稳定性。

# 结语
感谢团队的努力与支持。

汇报人：张三
日期：2026-01-10
"""
    out = HERE / "year-end-content.md"
    out.write_text(md, encoding="utf-8")
    return out


def main():
    HERE.mkdir(parents=True, exist_ok=True)
    print("生成:", make_year_end_template())
    print("生成:", make_realistic_body_style_template())
    print("生成:", make_gov_template())
    print("生成:", make_content_md())


if __name__ == "__main__":
    main()
