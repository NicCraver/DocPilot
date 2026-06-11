#!/usr/bin/env python3
"""Word 批量排版：根据 JSON 配置格式化 .docx 文件。"""
from __future__ import annotations

import json
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips

CITATION_RE = re.compile(r"[\[［]\d+(?:[，,\-\s]\d+)*[\]］]")
FRONT_MATTER_LINE = re.compile(
    r"^(摘要|关键词|中图分类号|文献标识码|文章编号|Abstract|Keywords|收稿日期|作者简介|引用格式|doi)",
    re.I,
)

SIZE_MAP: dict[str, float] = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
}


def builtin_default_config() -> dict[str, Any]:
    return {
        "page": {
            "margin_top": 2.54,
            "margin_bottom": 2.54,
            "margin_left": 3.17,
            "margin_right": 3.17,
            "footer_distance": 1.75,
            "page_number_align": "odd_even",
            "page_number_font": "宋体",
            "page_number_size": "四号",
            "force_a4": True,
        },
        "headings": {
            "title_font": "方正小标宋简体",
            "title_size": "二号",
            "title_line_spacing": 33,
            "subtitle_font": "楷体_GB2312",
            "subtitle_size": "三号",
            "subtitle_line_spacing": 33,
            "heading1_font": "黑体",
            "heading1_size": "三号",
            "heading2_font": "楷体_GB2312",
            "heading2_size": "三号",
            "body_font": "仿宋_GB2312",
            "body_size": "三号",
            "body_line_spacing": 28,
            "indent_left": 0,
            "indent_right": 0,
            "first_line_indent": 0,
        },
        "table": {
            "enabled": True,
            "auto_column_width": True,
            "unify_borders": True,
            "bold_header": True,
            "smart_align": False,
            "header_font": "仿宋_GB2312",
            "body_font": "仿宋_GB2312",
            "font_size": "小四",
            "row_spacing": 22,
            "row_height_cm": 0.7,
            "width_percent": 100,
            "border_pt": 0.5,
        },
        "other": {
            "table_caption_font": "黑体",
            "table_caption_size": "四号",
            "figure_caption_font": "黑体",
            "figure_caption_size": "四号",
            "attachment_enabled": True,
            "attachment_font": "黑体",
            "attachment_size": "三号",
            "auto_outline": True,
            "ascii_font": "Times New Roman",
            "enable_symbols": True,
            "collapse_empty_lines": True,
        },
    }


def thesis_default_config() -> dict[str, Any]:
    """学位论文常用排版预设（与前端 thesisWordTypesetConfig 对齐）。"""
    return {
        "page": {
            "margin_top": 2.5,
            "margin_bottom": 2.5,
            "margin_left": 3.0,
            "margin_right": 2.0,
            "footer_distance": 1.75,
            "page_number_align": "center",
            "page_number_font": "宋体",
            "page_number_size": "五号",
            "force_a4": True,
        },
        "headings": {
            "title_font": "黑体",
            "title_size": "二号",
            "title_line_spacing": 36,
            "subtitle_font": "宋体",
            "subtitle_size": "四号",
            "subtitle_line_spacing": 24,
            "heading1_font": "黑体",
            "heading1_size": "三号",
            "heading2_font": "黑体",
            "heading2_size": "四号",
            "body_font": "宋体",
            "body_size": "小四",
            "body_line_spacing": 22,
            "indent_left": 0,
            "indent_right": 0,
            "first_line_indent": 0.74,
        },
        "table": {
            "enabled": True,
            "auto_column_width": True,
            "unify_borders": True,
            "bold_header": True,
            "smart_align": True,
            "header_font": "宋体",
            "body_font": "宋体",
            "font_size": "五号",
            "row_spacing": 18,
            "row_height_cm": 0.6,
            "width_percent": 100,
            "border_pt": 0.5,
        },
        "other": {
            "table_caption_font": "宋体",
            "table_caption_size": "五号",
            "figure_caption_font": "宋体",
            "figure_caption_size": "五号",
            "attachment_enabled": False,
            "attachment_font": "宋体",
            "attachment_size": "小四",
            "auto_outline": True,
            "ascii_font": "Times New Roman",
            "enable_symbols": True,
            "collapse_empty_lines": True,
        },
    }


def journal_default_config() -> dict[str, Any]:
    """学术期刊排版预设（对齐煤炭工程等双栏论文）。"""
    return {
        "page": {
            "margin_top": 2.0,
            "margin_bottom": 2.0,
            "margin_left": 1.8,
            "margin_right": 1.6,
            "footer_distance": 1.5,
            "page_number_align": "left",
            "page_number_font": "宋体",
            "page_number_size": "五号",
            "force_a4": True,
            "columns": 2,
            "column_gap_cm": 0.75,
            "columns_start": "after_front_matter",
            "header_distance": 1.2,
            "journal_header": {
                "enabled": True,
                "font": "宋体",
                "size": "五号",
                "first_left": "第56卷第10期",
                "first_center": "煤炭工程",
                "first_right": "Vol. 56, No. 10",
                "running_left": "2024 年第10 期",
                "running_center": "煤炭工程",
                "running_right": "专家论坛",
                "center_char_spacing": True,
            },
        },
        "headings": {
            "title_font": "黑体",
            "title_size": "二号",
            "title_line_spacing": 28,
            "subtitle_font": "仿宋_GB2312",
            "subtitle_size": "小四",
            "subtitle_line_spacing": 18,
            "heading1_font": "黑体",
            "heading1_size": "小四",
            "heading2_font": "黑体",
            "heading2_size": "五号",
            "body_font": "宋体",
            "body_size": "五号",
            "body_line_spacing": 16,
            "indent_left": 0,
            "indent_right": 0,
            "first_line_indent": 0.64,
        },
        "table": {
            "enabled": True,
            "auto_column_width": True,
            "unify_borders": True,
            "bold_header": True,
            "smart_align": True,
            "header_font": "宋体",
            "body_font": "宋体",
            "font_size": "小五",
            "row_spacing": 14,
            "row_height_cm": 0.45,
            "width_percent": 100,
            "border_pt": 0.75,
        },
        "other": {
            "table_caption_font": "宋体",
            "table_caption_size": "五号",
            "figure_caption_font": "宋体",
            "figure_caption_size": "五号",
            "attachment_enabled": False,
            "attachment_font": "宋体",
            "attachment_size": "五号",
            "auto_outline": True,
            "ascii_font": "Times New Roman",
            "enable_symbols": True,
            "collapse_empty_lines": True,
            "abstract_font": "楷体_GB2312",
            "abstract_size": "五号",
            "abstract_hang_indent_cm": 0.64,
            "affiliation_font": "宋体",
            "affiliation_size": "五号",
            "three_line_table": True,
            "citation_superscript": True,
            "table_caption_en_size": "小五",
        },
    }


def resolve_pt(value: Any, default: float = 12) -> Pt:
    if isinstance(value, (int, float)):
        return Pt(float(value))
    if isinstance(value, str):
        text = value.strip()
        if text in SIZE_MAP:
            return Pt(SIZE_MAP[text])
        m = re.match(r"^([\d.]+)\s*pt$", text, re.I)
        if m:
            return Pt(float(m.group(1)))
    return Pt(default)


def ensure_rfonts(run, font_name: str, ascii_font: str | None = None) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    latin = ascii_font or font_name
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def style_run(
    run,
    font: str,
    size: Any,
    bold: bool | None = None,
    ascii_font: str | None = None,
) -> None:
    ensure_rfonts(run, font, ascii_font)
    run.font.size = resolve_pt(size)
    if bold is not None:
        run.font.bold = bold


def apply_line_spacing(paragraph, spacing_pt: float | None) -> None:
    if spacing_pt is None:
        return
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(spacing_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def paragraph_has_inline_drawing(paragraph) -> bool:
    """段落是否含 inline 插图（drawing/blip）。"""
    return bool(paragraph._element.xpath(".//w:drawing"))


def paragraph_style_name(paragraph) -> str:
    try:
        return (paragraph.style.name or "").strip()
    except Exception:
        return ""


def classify_paragraph(paragraph) -> str:
    name = paragraph_style_name(paragraph).lower()
    text = (paragraph.text or "").strip()

    # subtitle 必须在 title 之前：否则 "Subtitle" 会误匹配 "title" 子串
    exact_style_map = {
        "title": "title",
        "subtitle": "subtitle",
        "heading 1": "heading1",
        "heading 2": "heading2",
        "heading 3": "heading3",
        "标题": "title",
        "副标题": "subtitle",
        "标题 1": "heading1",
        "标题1": "heading1",
        "标题 2": "heading2",
        "标题2": "heading2",
        "标题 3": "heading3",
        "标题3": "heading3",
    }
    if name in exact_style_map:
        return exact_style_map[name]

    mapping = [
        ("subtitle", ("subtitle", "副标题")),
        ("title", ("title", "标题", "题目")),
        ("heading1", ("heading 1", "heading1", "标题 1", "标题1", "一级标题")),
        ("heading2", ("heading 2", "heading2", "标题 2", "标题2", "二级标题")),
        ("heading3", ("heading 3", "heading3", "标题 3", "标题3", "三级标题")),
        ("caption_table", ("table caption", "表题", "表格标题")),
        ("caption_figure", ("figure caption", "图题", "图形标题", "图片标题")),
        ("attachment", ("attachment", "附件")),
    ]
    for kind, keys in mapping:
        if any(k == name or (k in name and k != "title") for k in keys):
            return kind

    if re.match(r"^表\s*[\d一二三四五六七八九十]+", text):
        return "caption_table"
    if re.match(r"^Table\s+\d+", text, re.I):
        return "caption_table_en"
    if re.match(r"^Figure\s+\d+", text, re.I):
        return "caption_figure_en"
    if re.match(r"^doi[:：]", text, re.I):
        return "doi"
    if re.match(r"^图\s*[\d一二三四五六七八九十]+", text):
        return "caption_figure"
    if re.match(r"^附件\s*[\d一二三四五六七八九十]", text):
        return "attachment"
    if re.match(r"^摘要[:：]", text):
        return "abstract"
    if re.match(r"^关键词[:：]", text):
        return "keywords"
    if re.match(r"^中图分类号|^文献标识码|^文章编号", text):
        return "metadata"
    if re.match(r"^(Abstract|Keywords)[:：]?", text, re.I):
        return "english_block"
    if re.match(r"^[\(（][^)）]*[大学学院][^)）]*[\)）]", text):
        return "affiliation"
    if re.match(r"^(收稿日期|作者简介|引用格式)", text):
        return "footnote"
    if re.match(r"^\d+\.\s*\d+\s*[\u4e00-\u9fffA-Za-z]", text):
        return "heading2"
    if re.match(r"^\d{1,2}\s+[\u4e00-\u9fff]", text):
        return "heading1"
    if re.match(r"^[一二三四五六七八九十]+[、．.]", text):
        return "heading1"
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return "heading2"
    return "body"


def format_paragraph(paragraph, config: dict[str, Any], kind: str) -> None:
    if paragraph_has_inline_drawing(paragraph):
        # 含插图的段落不能设固定行距，否则 Word 只保留一行高度导致图片被裁切/遮挡
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        return

    headings = config["headings"]
    other = config["other"]

    spec: tuple[str, Any, float | None, bool | None] | None = None
    indent_left = headings.get("indent_left", 0)
    indent_right = headings.get("indent_right", 0)
    first_line_indent = headings.get("first_line_indent", 0)

    if kind == "title":
        spec = (
            headings["title_font"],
            headings["title_size"],
            headings.get("title_line_spacing"),
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "subtitle":
        spec = (
            headings["subtitle_font"],
            headings["subtitle_size"],
            headings.get("subtitle_line_spacing"),
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "heading1":
        spec = (headings["heading1_font"], headings["heading1_size"], None, True)
    elif kind == "heading2":
        text = (paragraph.text or "").strip()
        h2_bold = bool(re.match(r"^\d+\.", text))
        spec = (headings["heading2_font"], headings["heading2_size"], None, h2_bold)
    elif kind in ("heading3", "body"):
        spec = (
            headings["body_font"],
            headings["body_size"],
            headings.get("body_line_spacing"),
            None,
        )
    elif kind == "caption_table":
        spec = (
            other["table_caption_font"],
            other["table_caption_size"],
            None,
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "caption_figure":
        spec = (
            other["figure_caption_font"],
            other["figure_caption_size"],
            None,
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "caption_table_en":
        spec = (
            other.get("ascii_font") or "Times New Roman",
            other.get("table_caption_en_size") or "小五",
            None,
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "caption_figure_en":
        spec = (
            other.get("ascii_font") or "Times New Roman",
            other.get("table_caption_en_size") or "小五",
            None,
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "doi":
        spec = (
            headings["body_font"],
            "小五",
            None,
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif kind == "attachment" and other.get("attachment_enabled"):
        spec = (
            other["attachment_font"],
            other["attachment_size"],
            None,
            True,
        )
    elif kind == "abstract":
        spec = (
            other.get("abstract_font") or headings["body_font"],
            other.get("abstract_size") or headings["body_size"],
            headings.get("body_line_spacing"),
            None,
        )
        hang = float(other.get("abstract_hang_indent_cm") or 0)
        if hang > 0:
            pf = paragraph.paragraph_format
            pf.left_indent = Cm(hang)
            pf.first_line_indent = Cm(-hang)
    elif kind == "keywords":
        spec = (
            headings["heading1_font"],
            headings["body_size"],
            headings.get("body_line_spacing"),
            True,
        )
    elif kind == "metadata":
        spec = (
            headings["body_font"],
            headings["body_size"],
            headings.get("body_line_spacing"),
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "english_block":
        spec = (
            other.get("ascii_font") or "Times New Roman",
            headings["body_size"],
            headings.get("body_line_spacing"),
            None,
        )
    elif kind == "affiliation":
        spec = (
            other.get("affiliation_font") or headings["subtitle_font"],
            other.get("affiliation_size") or headings["subtitle_size"],
            None,
            None,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "footnote":
        spec = (
            headings["body_font"],
            "小五",
            None,
            None,
        )

    if spec is None:
        return

    font, size, line_spacing, bold = spec
    ascii_font = other.get("ascii_font") if other.get("ascii_font") else None
    if kind not in ("english_block",) and line_spacing is not None:
        apply_line_spacing(paragraph, line_spacing)
    elif kind == "english_block":
        apply_line_spacing(paragraph, line_spacing)
    pf = paragraph.paragraph_format
    if kind in ("body", "heading1", "heading2", "heading3", "attachment"):
        pf.left_indent = Cm(float(indent_left))
        pf.right_indent = Cm(float(indent_right))
        if float(first_line_indent) > 0:
            pf.first_line_indent = Cm(float(first_line_indent))
    elif kind == "abstract" and float(other.get("abstract_hang_indent_cm") or 0) <= 0:
        pf.left_indent = Cm(float(indent_left))
        pf.right_indent = Cm(float(indent_right))

    if paragraph.runs:
        for run in paragraph.runs:
            style_run(run, font, size, bold, ascii_font)
    elif paragraph.text:
        run = paragraph.add_run(paragraph.text)
        paragraph.text = ""
        style_run(run, font, size, bold, ascii_font)


def set_table_cell(
    cell,
    font: str,
    size: Any,
    bold: bool = False,
    ascii_font: str | None = None,
    row_spacing: float | None = None,
) -> None:
    for paragraph in cell.paragraphs:
        apply_line_spacing(paragraph, row_spacing)
        for run in paragraph.runs:
            style_run(run, font, size, bold, ascii_font)
        if not paragraph.runs and paragraph.text:
            run = paragraph.add_run(paragraph.text)
            paragraph.text = ""
            style_run(run, font, size, bold, ascii_font)


def format_tables(doc: Document, config: dict[str, Any]) -> None:
    table_cfg = config.get("table", {})
    if not table_cfg.get("enabled", True):
        return

    header_font = table_cfg.get("header_font", "黑体")
    body_font = table_cfg.get("body_font", "宋体")
    font_size = table_cfg.get("font_size", "五号")
    bold_header = table_cfg.get("bold_header", True)
    ascii_font = config.get("other", {}).get("ascii_font")
    border_pt = float(table_cfg.get("border_pt", 0.5))
    row_height_cm = float(table_cfg.get("row_height_cm", 0.6))
    row_spacing = float(table_cfg.get("row_spacing", 0) or 0)
    row_spacing_pt = row_spacing if row_spacing > 0 else None
    width_percent = int(table_cfg.get("width_percent", 100))
    auto_col = table_cfg.get("auto_column_width", True)

    for table in doc.tables:
        tbl_width: int | None = None
        if width_percent > 0:
            table.autofit = False
            try:
                tbl_width = int(Twips(9000 * width_percent / 100))
                table.width = tbl_width
            except Exception:
                pass

        if auto_col and tbl_width and len(table.columns) > 0:
            try:
                per_col = int(tbl_width / len(table.columns))
                for col in table.columns:
                    col.width = per_col
            except Exception:
                pass

        for row_idx, row in enumerate(table.rows):
            if row_height_cm > 0:
                try:
                    row.height = Cm(row_height_cm)
                except Exception:
                    pass
            for col_idx, cell in enumerate(row.cells):
                is_header = row_idx == 0 and bold_header
                font = header_font if is_header else body_font
                set_table_cell(
                    cell,
                    font,
                    font_size,
                    bold=is_header,
                    ascii_font=ascii_font,
                    row_spacing=row_spacing_pt,
                )
                if table_cfg.get("smart_align", True):
                    for p in cell.paragraphs:
                        if row_idx == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if table_cfg.get("unify_borders", True):
            if config.get("other", {}).get("three_line_table"):
                _apply_three_line_table_borders(table, border_pt)
            else:
                _apply_table_borders(table, border_pt)


def _apply_table_borders(table, border_pt: float) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    sz = str(int(border_pt * 8))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _create_page_number_field(run_element) -> None:
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_element.append(fld_char_begin)
    run_element.append(instr)
    run_element.append(fld_char_sep)
    run_element.append(fld_char_end)


def _set_section_columns(section, num: int, space_twips: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    if num <= 1:
        if qn("w:num") in cols.attrib:
            del cols.attrib[qn("w:num")]
        if qn("w:space") in cols.attrib:
            del cols.attrib[qn("w:space")]
        return
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def _copy_section_margins(src, dst) -> None:
    dst.top_margin = src.top_margin
    dst.bottom_margin = src.bottom_margin
    dst.left_margin = src.left_margin
    dst.right_margin = src.right_margin
    dst.footer_distance = src.footer_distance
    if src.page_width:
        dst.page_width = src.page_width
    if src.page_height:
        dst.page_height = src.page_height


def _insert_continuous_section_break(paragraph) -> None:
    """在段落后插入连续分节符（python-docx 的 add_break 不支持 section break）。"""
    p_pr = paragraph._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:sectPr"))
    if existing is not None:
        p_pr.remove(existing)
    sect_pr = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "continuous")
    sect_pr.append(sect_type)
    p_pr.append(sect_pr)


def apply_column_layout(doc: Document, config: dict[str, Any]) -> None:
    page = config.get("page", {})
    columns = int(page.get("columns", 1) or 1)
    if columns <= 1:
        return

    gap_cm = float(page.get("column_gap_cm", 0.75))
    space_twips = int(gap_cm / 2.54 * 1440)
    start_mode = page.get("columns_start", "first_numeric_heading")

    target_idx = _find_column_break_index(doc, start_mode)

    if target_idx is None or target_idx == 0:
        _set_section_columns(doc.sections[0], columns, space_twips)
        return

    break_paragraph = doc.paragraphs[target_idx - 1]
    _insert_continuous_section_break(break_paragraph)
    if len(doc.sections) > 1:
        new_section = doc.sections[-1]
        _copy_section_margins(doc.sections[0], new_section)
        _set_section_columns(new_section, columns, space_twips)


def _apply_three_line_table_borders(table, border_pt: float) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    thick = str(int(border_pt * 8 * 1.5))
    thin = str(int(border_pt * 8))
    edge_specs = {
        "top": thick,
        "bottom": thick,
        "insideH": thin,
        "left": "0",
        "right": "0",
        "insideV": "0",
    }
    for edge, sz in edge_specs.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if sz == "0":
            element.set(qn("w:val"), "none")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), sz)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")


def _space_journal_center(text: str, spaced: bool) -> str:
    if not spaced or not text:
        return text
    compact = text.replace(" ", "")
    if re.fullmatch(r"[\u4e00-\u9fff]+", compact):
        return " ".join(compact)
    return text


def _clear_paragraph_runs(paragraph) -> None:
    element = paragraph._element
    for child in list(element.findall(qn("w:r"))):
        element.remove(child)


def _fill_triple_header(header, left: str, center: str, right: str, hdr_cfg: dict[str, Any]) -> None:
    font = hdr_cfg.get("font", "宋体")
    size = hdr_cfg.get("size", "五号")
    ascii_font = hdr_cfg.get("ascii_font") or "Times New Roman"
    spaced = hdr_cfg.get("center_char_spacing", True)
    center = _space_journal_center(center, spaced)

    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    _clear_paragraph_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(Cm(8.4), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT)

    run = paragraph.add_run(left or "")
    style_run(run, font, size, None, font)
    last_run = run
    if center:
        last_run.add_tab()
        r_center = paragraph.add_run(center)
        style_run(r_center, font, size, True, font)
        last_run = r_center
    if right:
        last_run.add_tab()
        r_right = paragraph.add_run(right)
        latin = ascii_font if re.search(r"[A-Za-z]", right) else font
        style_run(r_right, font, size, None, latin)


def apply_journal_headers(doc: Document, config: dict[str, Any]) -> None:
    page = config.get("page", {})
    hdr = page.get("journal_header") or {}
    if not hdr.get("enabled"):
        return

    hdr = {
        **hdr,
        "ascii_font": config.get("other", {}).get("ascii_font", "Times New Roman"),
    }
    header_distance = float(page.get("header_distance", 1.2))

    sections = doc.sections
    if not sections:
        return

    first = hdr.get("first_left") or hdr.get("first_center") or hdr.get("first_right")
    running = hdr.get("running_left") or hdr.get("running_center") or hdr.get("running_right")

    s0 = sections[0]
    s0.header_distance = Cm(header_distance)
    if first:
        s0.different_first_page_header_footer = True
        _fill_triple_header(
            s0.first_page_header,
            hdr.get("first_left", ""),
            hdr.get("first_center", ""),
            hdr.get("first_right", ""),
            hdr,
        )
        if s0.header.paragraphs:
            _clear_paragraph_runs(s0.header.paragraphs[0])
        else:
            s0.header.add_paragraph()

    start = 1 if first else 0
    for section in sections[start:]:
        section.header_distance = Cm(header_distance)
        if running:
            _fill_triple_header(
                section.header,
                hdr.get("running_left", ""),
                hdr.get("running_center", ""),
                hdr.get("running_right", ""),
                hdr,
            )


def _find_column_break_index(doc: Document, start_mode: str) -> int | None:
    if start_mode in ("never", None, ""):
        return None
    if start_mode == "first_numeric_heading":
        for i, paragraph in enumerate(doc.paragraphs):
            text = (paragraph.text or "").strip()
            if re.match(r"^\d{1,2}\s+[\u4e00-\u9fff]", text):
                return i
        return None
    if start_mode == "after_front_matter":
        last_front: int | None = None
        for i, paragraph in enumerate(doc.paragraphs):
            text = (paragraph.text or "").strip()
            if not text:
                continue
            if FRONT_MATTER_LINE.match(text):
                last_front = i
                continue
            if re.match(r"^\d{1,2}\s+[\u4e00-\u9fff]", text):
                return i
        if last_front is not None:
            return last_front + 1
        return None
    return None


def _apply_citations_to_paragraph(
    paragraph,
    font: str,
    size: Any,
    ascii_font: str | None,
) -> None:
    if paragraph_has_inline_drawing(paragraph):
        return
    full = paragraph.text or ""
    if not CITATION_RE.search(full):
        return

    parts: list[tuple[str, str]] = []
    pos = 0
    for match in CITATION_RE.finditer(full):
        if match.start() > pos:
            parts.append(("text", full[pos : match.start()]))
        parts.append(("cite", match.group()))
        pos = match.end()
    if pos < len(full):
        parts.append(("text", full[pos:]))

    _clear_paragraph_runs(paragraph)
    for kind, chunk in parts:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if kind == "cite":
            run.font.superscript = True
        style_run(run, font, size, None, ascii_font)


def apply_citation_superscripts(doc: Document, config: dict[str, Any]) -> None:
    if not config.get("other", {}).get("citation_superscript"):
        return
    headings = config["headings"]
    other = config["other"]
    font = headings["body_font"]
    size = headings["body_size"]
    ascii_font = other.get("ascii_font")

    skip = {
        "caption_table",
        "caption_table_en",
        "caption_figure",
        "caption_figure_en",
        "title",
        "subtitle",
        "footnote",
        "metadata",
        "doi",
        "affiliation",
    }
    for paragraph in doc.paragraphs:
        if classify_paragraph(paragraph) in skip:
            continue
        _apply_citations_to_paragraph(paragraph, font, size, ascii_font)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_citations_to_paragraph(paragraph, font, size, ascii_font)


def _add_footer_page_number(section, align: str, font: str, size: Any) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()

    if align == "odd_even":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    ensure_rfonts(run, font, font)
    run.font.size = resolve_pt(size)
    r = run._element
    _create_page_number_field(r)


def apply_page_setup(doc: Document, config: dict[str, Any]) -> None:
    page = config.get("page", {})
    for section in doc.sections:
        section.top_margin = Cm(float(page.get("margin_top", 2.54)))
        section.bottom_margin = Cm(float(page.get("margin_bottom", 2.54)))
        section.left_margin = Cm(float(page.get("margin_left", 3.17)))
        section.right_margin = Cm(float(page.get("margin_right", 3.17)))
        section.footer_distance = Cm(float(page.get("footer_distance", 1.75)))
        if page.get("force_a4", True):
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)


def apply_page_footers(doc: Document, config: dict[str, Any]) -> None:
    page = config.get("page", {})
    for section in doc.sections:
        _add_footer_page_number(
            section,
            page.get("page_number_align", "odd_even"),
            page.get("page_number_font", "宋体"),
            page.get("page_number_size", "四号"),
        )


def apply_outline_levels(doc: Document, config: dict[str, Any]) -> None:
    if not config.get("other", {}).get("auto_outline", True):
        return
    for paragraph in doc.paragraphs:
        kind = classify_paragraph(paragraph)
        level_map = {
            "heading1": "0",
            "heading2": "1",
            "heading3": "2",
            "title": "0",
        }
        lvl = level_map.get(kind)
        if lvl is None:
            continue
        p_pr = paragraph._element.get_or_add_pPr()
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            p_pr.append(outline)
        outline.set(qn("w:val"), lvl)


def normalize_symbols(text: str) -> str:
    replacements = {
        "•": "·",
        "●": "·",
        "○": "○",
        "▪": "·",
        "■": "·",
        "–": "-",
        "—": "-",
        "…": "...",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text


def apply_symbol_normalization(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                run.text = normalize_symbols(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = normalize_symbols(run.text)


def infer_paragraph_style(line: str) -> str | None:
    text = line.strip()
    if not text:
        return None
    if text.startswith("### "):
        return "Heading 3"
    if text.startswith("## "):
        return "Heading 2"
    if text.startswith("# "):
        return "Heading 1"
    if re.match(r"^[一二三四五六七八九十]+[、．.]", text):
        return "Heading 1"
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return "Heading 2"
    return "Normal"


def build_document_from_text(text: str, config: dict[str, Any]) -> Document:
    other = config.get("other", {})
    collapse = other.get("collapse_empty_lines", True)
    doc = Document()
    prev_empty = False
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            if collapse:
                if prev_empty:
                    continue
                prev_empty = True
            doc.add_paragraph("")
            continue
        prev_empty = False
        style = infer_paragraph_style(line)
        content = line
        if style and style.startswith("Heading") and content.startswith("#"):
            content = re.sub(r"^#+\s*", "", content)
        doc.add_paragraph(content, style=style if style else "Normal")
    return doc


def format_document(doc: Document, config: dict[str, Any]) -> None:
    apply_page_setup(doc, config)
    apply_column_layout(doc, config)
    apply_page_setup(doc, config)
    apply_page_footers(doc, config)
    apply_journal_headers(doc, config)

    other = config.get("other", {})
    if other.get("enable_symbols"):
        apply_symbol_normalization(doc)

    for paragraph in doc.paragraphs:
        kind = classify_paragraph(paragraph)
        if kind == "attachment" and not other.get("attachment_enabled"):
            kind = "body"
        format_paragraph(paragraph, config, kind)

    format_tables(doc, config)
    apply_citation_superscripts(doc, config)
    apply_outline_levels(doc, config)


def merge_config(user: dict[str, Any] | None) -> dict[str, Any]:
    base = builtin_default_config()
    if not user:
        return base
    result = deepcopy(base)

    def deep_merge(dst: dict, src: dict) -> None:
        for key, value in src.items():
            if isinstance(value, dict) and isinstance(dst.get(key), dict):
                deep_merge(dst[key], value)
            else:
                dst[key] = value

    deep_merge(result, user)
    return result


def resolve_batch_output_path(
    src: Path,
    output_mode: str,
    output_dir: str | None,
    output_suffix: str,
) -> Path:
    if output_mode == "output_dir":
        if not output_dir:
            raise ValueError("输出到文件夹模式需要指定 output_dir")
        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        return out_dir / src.name
    if output_mode == "suffix":
        suffix = output_suffix or "_排版"
        stem = src.stem
        if stem.endswith(suffix):
            return src.with_name(f"{stem}.docx")
        return src.with_name(f"{stem}{suffix}.docx")
    return src


def process_file(
    input_path: str,
    config: dict[str, Any],
    output_mode: str = "in_place",
    output_dir: str | None = None,
    output_suffix: str = "_排版",
) -> dict[str, Any]:
    src = Path(input_path)
    if not src.is_file():
        raise FileNotFoundError(f"文件不存在: {input_path}")
    if src.name.startswith("~$"):
        raise ValueError(f"跳过 Word 临时文件: {input_path}")
    if src.suffix.lower() != ".docx":
        raise ValueError(f"仅支持 .docx: {input_path}")

    doc = Document(str(src))
    format_document(doc, config)

    out_path = resolve_batch_output_path(src, output_mode, output_dir, output_suffix)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    result: dict[str, Any] = {"input": str(src), "output": str(out_path), "ok": True}
    if output_mode == "in_place":
        backup = src.with_suffix(src.suffix + ".bak")
        shutil.copy2(src, backup)
        doc.save(str(out_path))
        result["backup"] = str(backup)
    else:
        doc.save(str(out_path))
    return result


def process_text(text: str, output_path: str, config: dict[str, Any]) -> str:
    doc = build_document_from_text(text, config)
    format_document(doc, config)
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


def main() -> int:
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "缺少 JSON 参数"}, ensure_ascii=False))
        return 1

    try:
        payload = json.loads(sys.argv[1])
    except json.JSONDecodeError as exc:
        print(json.dumps({"ok": False, "error": f"JSON 解析失败: {exc}"}, ensure_ascii=False))
        return 1

    mode = payload.get("mode", "batch")
    config = merge_config(payload.get("config"))
    logs: list[str] = []
    results: list[dict[str, str]] = []

    try:
        if mode == "text":
            text = payload.get("text", "")
            output_path = payload.get("output_path", "")
            if not output_path:
                raise ValueError("缺少 output_path")
            saved = process_text(text, output_path, config)
            logs.append(f"已从文本生成并排版: {saved}")
            results.append({"input": "<text>", "output": saved})
        else:
            paths = payload.get("input_paths") or []
            if not paths:
                raise ValueError("input_paths 为空")
            output_mode = payload.get("output_mode")
            legacy_output = payload.get("output_path")
            if not output_mode:
                if bool(payload.get("in_place", True)):
                    output_mode = "in_place"
                elif legacy_output:
                    output_mode = "explicit"
                else:
                    output_mode = "suffix"
            output_dir = payload.get("output_dir")
            output_suffix = str(payload.get("output_suffix") or "_排版")
            continue_on_error = bool(payload.get("continue_on_error", True))
            had_error = False
            for raw in paths:
                inp = str(raw)
                try:
                    if output_mode == "explicit" and legacy_output:
                        src = Path(inp)
                        if not src.is_file():
                            raise FileNotFoundError(f"文件不存在: {inp}")
                        doc = Document(str(src))
                        format_document(doc, config)
                        out_path = Path(str(legacy_output))
                        out_path.parent.mkdir(parents=True, exist_ok=True)
                        doc.save(str(out_path))
                        item = {"input": inp, "output": str(out_path), "ok": True}
                    else:
                        item = process_file(
                            inp,
                            config,
                            output_mode=str(output_mode),
                            output_dir=str(output_dir) if output_dir else None,
                            output_suffix=output_suffix,
                        )
                    backup_note = ""
                    if item.get("backup"):
                        backup_note = f"（备份: {item['backup']}）"
                    logs.append(f"排版完成: {item['output']}{backup_note}")
                    results.append(item)
                except Exception as exc:
                    had_error = True
                    err_item = {"input": inp, "output": "", "ok": False, "error": str(exc)}
                    results.append(err_item)
                    logs.append(f"排版失败: {inp} — {exc}")
                    if not continue_on_error:
                        raise
            if had_error and continue_on_error:
                ok_count = sum(1 for r in results if r.get("ok"))
                logs.append(f"批量结束: 成功 {ok_count}/{len(results)}")

        print(
            json.dumps(
                {"ok": True, "results": results, "logs": logs, "config": config},
                ensure_ascii=False,
            )
        )
        return 0
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc), "logs": logs}, ensure_ascii=False))
        return 1


if __name__ == "__main__":
    sys.exit(main())
