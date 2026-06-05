#!/usr/bin/env python3
"""生成约1000字学术论文测试稿（含配图），并用论文格式预设执行排版与校验。"""
from __future__ import annotations

import importlib.util
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

# 与 word-typeset.py 保持一致，用于校验
def paragraph_has_inline_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))

ROOT = Path(__file__).resolve().parent.parent
SCRIPT = ROOT / "scripts" / "word-typeset.py"
DATA = ROOT / "scripts" / "word-typeset-test-data"
ASSETS = DATA / "assets"
TEXT_SRC = DATA / "学术论文-测试原文.txt"
DOC_IN = DATA / "学术论文-排版前.docx"
DOC_OUT = DATA / "学术论文-排版后.docx"
CHART_PNG = ASSETS / "thesis-framework-chart.png"
BAR_PNG = ASSETS / "thesis-performance-chart.png"

FONT_ALIASES: dict[str, set[str]] = {
    "宋体": {"宋体", "SimSun", "simsun", "Songti SC", "STSong"},
    "黑体": {"黑体", "SimHei", "simhei", "Heiti SC", "STHeiti"},
}


def run_east_asia_font(run) -> str | None:
    r_pr = run._element.rPr
    if r_pr is None:
        return run.font.name
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        return run.font.name
    return r_fonts.get(qn("w:eastAsia")) or run.font.name


def font_matches(actual: str | None, expected: str) -> bool:
    act = (actual or "").strip()
    if not act:
        return False
    aliases = FONT_ALIASES.get(expected, {expected})
    return act in aliases or act.lower() in {a.lower() for a in aliases}


def load_module():
    spec = importlib.util.spec_from_file_location("word_typeset", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


def char_count(text: str) -> int:
    return len("".join(ch for ch in text if not ch.isspace()))


def find_cjk_font(size: int = 16) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    for path in candidates:
        if Path(path).is_file():
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                continue
    return ImageFont.load_default()


def _draw_round_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    fill: tuple[int, int, int],
    outline: tuple[int, int, int],
    radius: int = 8,
) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2)


def _draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: tuple[int, int, int],
) -> None:
    draw.line([start, end], fill=color, width=2)
    ex, ey = end
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        tip = (-8, -4) if end[0] > start[0] else (8, -4)
    else:
        tip = (-4, -8) if end[1] > start[1] else (-4, 8)
    draw.polygon([(ex, ey), (ex + tip[0], ey + tip[1]), (ex + tip[1], ey + tip[0])], fill=color)


def _text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font) -> None:
    x0, y0, x1, y1 = box
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text((x0 + (x1 - x0 - tw) / 2, y0 + (y1 - y0 - th) / 2), text, fill=(33, 37, 41), font=font)


def draw_framework_diagram(path: Path) -> None:
    """智慧社区治理框架示意图：三模块 + 底层治理绩效。"""
    w, h = 720, 420
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_font = find_cjk_font(22)
    box_font = find_cjk_font(18)
    sub_font = find_cjk_font(14)

    draw.text((36, 24), "智慧社区治理框架示意图", fill=(30, 64, 175), font=title_font)

    boxes = [
        (48, 120, 228, 220, "平台赋能", (219, 234, 254), (59, 130, 246)),
        (270, 120, 450, 220, "多元协同", (220, 252, 231), (34, 197, 94)),
        (492, 120, 672, 220, "数据驱动", (254, 243, 199), (245, 158, 11)),
    ]
    for x0, y0, x1, y1, label, fill, outline in boxes:
        _draw_round_box(draw, (x0, y0, x1, y1), fill, outline)
        _text_center(draw, (x0, y0, x1, y1), label, box_font)

    _draw_arrow(draw, (228, 170), (270, 170), (100, 116, 139))
    _draw_arrow(draw, (450, 170), (492, 170), (100, 116, 139))
    _draw_arrow(draw, (360, 220), (360, 280), (148, 163, 184))

    _draw_round_box(draw, (120, 280, 600, 360), (241, 245, 249), (148, 163, 184), radius=12)
    _text_center(draw, (120, 280, 600, 360), "治理绩效：响应效率 · 居民参与 · 满意度", sub_font)

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, format="PNG")


def draw_performance_chart(path: Path) -> None:
    """三社区治理绩效对比柱状图。"""
    w, h = 720, 420
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = find_cjk_font(22)
    label_font = find_cjk_font(16)
    value_font = find_cjk_font(14)

    draw.text((36, 24), "三社区治理绩效对比", fill=(30, 64, 175), font=title_font)

    chart = (80, 80, 660, 340)
    x0, y0, x1, y1 = chart
    draw.rectangle(chart, outline=(203, 213, 225), width=1)
    draw.line([(x0, y1), (x1, y1)], fill=(100, 116, 139), width=2)

    items = [
        ("A 社区", 82, (59, 130, 246)),
        ("B 社区", 68, (34, 197, 94)),
        ("C 社区", 74, (245, 158, 11)),
    ]
    bar_w = 90
    gap = 70
    base_x = x0 + 80
    max_h = y1 - y0 - 50
    for i, (label, score, color) in enumerate(items):
        bx = base_x + i * (bar_w + gap)
        bh = int(max_h * score / 100)
        by = y1 - bh
        draw.rectangle((bx, by, bx + bar_w, y1), fill=color, outline=color)
        draw.text((bx + 18, y1 + 8), label, fill=(51, 65, 85), font=label_font)
        draw.text((bx + 24, by - 22), f"{score}%", fill=(51, 65, 85), font=value_font)

    draw.text((x0 + 8, y0 + 12), "绩效指数", fill=(100, 116, 139), font=value_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, format="PNG")


def ensure_assets() -> None:
    draw_framework_diagram(CHART_PNG)
    draw_performance_chart(BAR_PNG)


def build_docx_from_text(text: str, path: Path) -> None:
    """按学位论文结构生成未排版 docx（含表格与插图）。"""
    ensure_assets()
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    doc = Document()
    if not lines:
        raise ValueError("原文为空")

    doc.add_paragraph(lines[0], style="Title")
    if len(lines) > 1:
        doc.add_paragraph(lines[1], style="Subtitle")

    figure_idx = 0
    figures = [CHART_PNG, BAR_PNG]
    i = 2
    while i < len(lines):
        line = lines[i]
        if line.startswith("表") and re.match(r"^表\s*[\d一二三四五六七八九十]+", line):
            doc.add_paragraph(line)
            table = doc.add_table(rows=4, cols=4)
            hdr = ["模式类型", "技术特征", "协同机制", "适用场景"]
            for c, h in enumerate(hdr):
                table.rows[0].cells[c].text = h
            rows_data = [
                ("平台主导型", "统一数据中台", "政府牵头", "老旧改造社区"),
                ("市场运营型", "SaaS 服务", "物业主导", "新建商品房"),
                ("混合共建型", "模块化部署", "多方共治", "城乡结合部"),
            ]
            for r, row in enumerate(rows_data, 1):
                for c, val in enumerate(row):
                    table.rows[r].cells[c].text = val
            i += 1
            continue
        if line.startswith("图"):
            if figure_idx < len(figures):
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pic.add_run()
                run.add_picture(str(figures[figure_idx]), width=Inches(4.8))
                figure_idx += 1
            doc.add_paragraph(line)
            i += 1
            continue
        if line.startswith("参考文献"):
            doc.add_paragraph(line, style="Heading 1")
            i += 1
            continue
        if line.startswith("关键词") or line == "摘要":
            doc.add_paragraph(line, style="Heading 1")
        elif line.startswith("（"):
            doc.add_paragraph(line, style="Heading 2")
        elif line.startswith(("一、", "二、", "三、", "四、", "五、")):
            doc.add_paragraph(line, style="Heading 1")
        elif line.startswith("["):
            doc.add_paragraph(line, style="Normal")
        else:
            doc.add_paragraph(line, style="Normal")
        i += 1

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def run_typeset(inp: Path, out: Path, config: dict) -> None:
    py = ROOT / ".venv" / "bin" / "python3"
    if not py.is_file():
        py = Path(sys.executable)
    payload = {
        "mode": "batch",
        "input_paths": [str(inp)],
        "output_path": str(out),
        "in_place": False,
        "config": config,
    }
    proc = subprocess.run(
        [str(py), str(SCRIPT), json.dumps(payload, ensure_ascii=False)],
        capture_output=True,
        text=True,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr or proc.stdout)
    result = json.loads(proc.stdout.strip())
    if not result.get("ok"):
        raise RuntimeError(result.get("error", "排版失败"))


def verify_thesis_output(doc_path: Path, config: dict) -> list[tuple[str, bool, str]]:
    doc = Document(str(doc_path))
    body_font = config["headings"]["body_font"]
    body_size = config["headings"]["body_size"]
    checks: list[tuple[str, bool, str]] = []

    body_paras = [
        p
        for p in doc.paragraphs
        if p.text.strip()
        and not p.text.startswith(("图", "表", "一、", "二、", "三、", "四、", "五、"))
        and not p.text.startswith("（")
        and p.text not in ("摘要", "关键词：智慧城市；社区治理；数字平台；协同治理", "参考文献")
        and len(p.text) > 20
    ]
    if body_paras:
        run = body_paras[0].runs[0] if body_paras[0].runs else None
        east = run_east_asia_font(run) if run else None
        font_ok = font_matches(east, body_font)
        checks.append(("正文宋体", font_ok, east or "无 run"))
    else:
        checks.append(("正文宋体", False, "未找到正文段"))

    h1 = next((p for p in doc.paragraphs if p.text.startswith("一、")), None)
    if h1 and h1.runs:
        east = run_east_asia_font(h1.runs[0])
        checks.append(
            ("一级标题黑体", font_matches(east, config["headings"]["heading1_font"]), east or "")
        )
    else:
        checks.append(("一级标题黑体", False, "未找到一级标题"))

    fig_caps = [p for p in doc.paragraphs if p.text.startswith("图")]
    checks.append(("图题保留", len(fig_caps) >= 2, f"图题 {len(fig_caps)} 条"))

    inline_shapes = sum(len(p._element.xpath(".//wp:inline")) for p in doc.paragraphs)
    checks.append(("插图保留", inline_shapes >= 2, f"inline {inline_shapes} 个"))

    image_paras = [p for p in doc.paragraphs if paragraph_has_inline_drawing(p)]
    if image_paras:
        bad = [
            p
            for p in image_paras
            if p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
        ]
        checks.append(
            (
                "插图段落非固定行距",
                len(bad) == 0,
                f"{len(image_paras) - len(bad)}/{len(image_paras)} 段正常",
            )
        )
    else:
        checks.append(("插图段落非固定行距", False, "未找到插图段"))

    sect = doc.sections[0]
    left_cm = round(sect.left_margin.cm, 2)
    checks.append(
        (
            "左边距约3cm",
            abs(left_cm - config["page"]["margin_left"]) < 0.2,
            f"{left_cm} cm",
        )
    )

    checks.append(("表格保留", len(doc.tables) >= 1, f"{len(doc.tables)} 张"))
    return checks


def main() -> int:
    wt = load_module()
    if not hasattr(wt, "thesis_default_config"):
        raise RuntimeError("word-typeset.py 缺少 thesis_default_config()")
    text = TEXT_SRC.read_text(encoding="utf-8")
    n = char_count(text)
    print(f"==> 学术论文测试稿（汉字约 {n} 字）")
    print(f"    原文: {TEXT_SRC}")

    build_docx_from_text(text, DOC_IN)
    print(f"    生成: {DOC_IN}（含表格 1 张、插图 2 张）")

    if DOC_OUT.exists():
        DOC_OUT.unlink()
    config = wt.thesis_default_config()
    run_typeset(DOC_IN, DOC_OUT, config)
    print(f"    排版: {DOC_OUT}")

    doc = Document(str(DOC_OUT))
    print(f"    段落 {len(doc.paragraphs)} 段，表格 {len(doc.tables)} 张")

    checks = verify_thesis_output(DOC_OUT, config)
    failed = 0
    print("\n==> 论文格式校验")
    for name, ok, detail in checks:
        mark = "✓" if ok else "✗"
        print(f"    {mark} {name}: {detail}")
        if not ok:
            failed += 1

    if failed:
        print(f"\n失败 {failed} 项")
        return 1
    print("\n完成。可用 Word 打开「学术论文-排版后.docx」查看效果。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
