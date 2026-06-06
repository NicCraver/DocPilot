#!/usr/bin/env python3
"""
按 Word 模板生成文档：保留模板版式，将用户内容填入对应章节，并移除红色填写说明。
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

# 小四 = 12pt；模板要求正文 1.5 倍行距
BODY_FONT = "宋体"
BODY_SIZE_PT = 12

SECTION_DEFS: list[dict[str, Any]] = [
    {"key": "preface", "labels": ["前言"], "heading_hints": ["前言"]},
    {
        "key": "achievements",
        "labels": ["2.1", "工作成果", "核心贡献"],
        "heading_hints": ["2.1", "工作成果"],
    },
    {
        "key": "review",
        "labels": ["2.2", "工作复盘", "反思"],
        "heading_hints": ["2.2", "工作复盘"],
    },
    {
        "key": "competency",
        "labels": ["个人岗位胜任", "能力评估", "胜任度"],
        "heading_hints": ["个人岗位胜任", "能力评估"],
    },
    {
        "key": "plan2026",
        "labels": ["五要点", "工作计划", "2026年度个人"],
        "heading_hints": ["五要点", "工作计划"],
    },
    {
        "key": "outlook",
        "labels": ["未来展望", "未来展望与规划"],
        "heading_hints": ["未来展望"],
    },
    {"key": "conclusion", "labels": ["结语"], "heading_hints": ["结语"]},
]

GUIDE_KEYWORDS = (
    "填写指南",
    "请点击链接",
    "扫描二维码",
    "feishu.cn",
    "标题格式（使用时用格式刷即可）",
    "二级标题/内容",
    "三级标题/内容",
    "四级标题/内容",
    "五级标题/内容",
    "适用于美腾科技",
    "文档中出现的红色斜字体",
    "年终总结报告中涉及填写指导",
    "正文填写内容统一字体",
)


def log_lines() -> list[str]:
    return []


def _normalize(text: str) -> str:
    return re.sub(r"\s+", "", text.strip())


def _run_is_red_instruction(run) -> bool:
    color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
    if color == RGBColor(0xFF, 0x00, 0x00):
        return True
    if run.italic and color == RGBColor(0xFF, 0x00, 0x00):
        return True
    return False


def paragraph_is_red_instruction(para: Paragraph) -> bool:
    text = para.text.strip()
    if not text:
        return False
    red_runs = sum(1 for r in para.runs if _run_is_red_instruction(r))
    if red_runs == 0:
        return False
    # 红色斜体说明段，或整段以（ 开头且含红色
    if text.startswith("（") and red_runs > 0:
        return True
    if red_runs >= max(1, len(para.runs) // 2):
        return True
    return False


def paragraph_is_guide(para: Paragraph) -> bool:
    text = para.text.strip()
    if not text:
        return False
    if paragraph_is_red_instruction(para):
        return True
    return any(k in text for k in GUIDE_KEYWORDS)


def match_section_key(text: str, *, heading_mode: bool = False) -> str | None:
    norm = _normalize(text)
    if not norm:
        return None
    if heading_mode and len(norm) > 80:
        return None
    for item in SECTION_DEFS:
        hints = item["heading_hints"] if heading_mode else item["labels"]
        for hint in hints:
            hint_norm = _normalize(hint)
            if heading_mode:
                if norm == hint_norm or norm.startswith(hint_norm) or hint_norm in norm[:40]:
                    return item["key"]
            elif hint_norm in norm:
                return item["key"]
    if heading_mode and "汇报人" in text:
        return "reporter"
    if heading_mode and ("日期：" in text) and len(text) < 40:
        return "date"
    return None


def is_template_heading(para: Paragraph) -> str | None:
    text = para.text.strip()
    if not text:
        return None
    if paragraph_is_red_instruction(para) or paragraph_is_guide(para):
        return None
    key = match_section_key(text, heading_mode=True)
    if not key or key in {"reporter", "date"}:
        if "汇报人" in text:
            return "reporter"
        if "日期：" in text and len(text) < 40:
            return "date"
        return None
    bold_len = sum(len(r.text) for r in para.runs if r.bold)
    numbered = bool(re.match(r"^[\d\.]+", text))
    if bold_len >= max(2, len(text) * 0.4) or numbered or len(text) <= 20:
        return key
    return None


def split_paragraphs(text: str) -> list[str]:
    chunks = re.split(r"\n\s*\n+", text.replace("\r\n", "\n"))
    out: list[str] = []
    for chunk in chunks:
        line = chunk.strip()
        if line:
            out.append(line)
    return out


def parse_markdown_sections(raw: str) -> dict[str, list[str]]:
    lines = raw.replace("\r\n", "\n").split("\n")
    current_key: str | None = None
    buffer: list[str] = []
    sections: dict[str, list[str]] = {d["key"]: [] for d in SECTION_DEFS}
    meta: dict[str, str] = {}

    def flush() -> None:
        nonlocal buffer, current_key
        if current_key and buffer:
            sections.setdefault(current_key, []).extend(split_paragraphs("\n".join(buffer)))
        buffer = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                buffer.append("")
            continue

        heading = re.match(r"^#{1,4}\s+(.+)$", stripped)
        if heading:
            flush()
            new_key = match_section_key(heading.group(1), heading_mode=True)
            if new_key:
                current_key = new_key
            # 子标题（如 ## 重点项目）不切换章节，内容仍归入当前节
            continue

        m_reporter = re.match(r"汇报人[：:]\s*(.+)$", stripped)
        m_date = re.match(r"日期[：:]\s*(.+)$", stripped)
        if m_reporter:
            meta["reporter"] = m_reporter.group(1).strip()
            continue
        if m_date:
            meta["date"] = m_date.group(1).strip()
            continue

        key = match_section_key(stripped, heading_mode=True)
        if key and key not in {"reporter", "date"} and len(stripped) < 80:
            flush()
            current_key = key
            continue

        if current_key:
            buffer.append(stripped)
        else:
            # 无标题前缀时归入前言
            current_key = "preface"
            buffer.append(stripped)

    flush()
    sections["_meta"] = [json.dumps(meta, ensure_ascii=False)]
    return sections


def parse_plain_text(raw: str) -> dict[str, list[str]]:
    return parse_markdown_sections(raw)


def parse_docx_content(path: Path) -> dict[str, list[str]]:
    doc = Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if lines and lines[-1] != "":
                lines.append("")
            continue
        if paragraph_is_guide(para) or paragraph_is_red_instruction(para):
            continue
        lines.append(text)
    return parse_markdown_sections("\n".join(lines))


def load_sections(content_path: str | None, content_text: str | None, content_kind: str) -> dict[str, list[str]]:
    if content_text and content_text.strip():
        raw = content_text
    elif content_path:
        path = Path(content_path)
        ext = path.suffix.lower()
        if ext == ".docx":
            return parse_docx_content(path)
        raw = path.read_text(encoding="utf-8")
    else:
        raise ValueError("请提供内容文件或文本")

    if content_kind == "markdown" or (content_path and Path(content_path).suffix.lower() in {".md", ".markdown"}):
        return parse_markdown_sections(raw)
    return parse_plain_text(raw)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _set_run_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", font_name)


def apply_body_style(para: Paragraph) -> None:
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(24)  # 约 2 字符首行缩进
    for run in para.runs:
        _set_run_east_asia_font(run, BODY_FONT)
        run.font.size = Pt(BODY_SIZE_PT)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.italic = False
        run.bold = False


def set_paragraph_text(para: Paragraph, text: str) -> None:
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    run = para.add_run(text)
    _set_run_east_asia_font(run, BODY_FONT)
    run.font.size = Pt(BODY_SIZE_PT)
    run.font.color.rgb = RGBColor(0, 0, 0)
    apply_body_style(para)


def fill_template(
    template_path: str,
    output_path: str,
    sections: dict[str, list[str]],
    reporter: str | None = None,
    report_date: str | None = None,
) -> list[str]:
    logs: list[str] = []
    src = Path(template_path)
    dst = Path(output_path)
    if not src.is_file():
        raise FileNotFoundError(f"模板不存在: {template_path}")

    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    doc = Document(str(dst))
    paragraphs = list(doc.paragraphs)

    meta_raw = sections.get("_meta", [])
    if meta_raw:
        try:
            meta = json.loads(meta_raw[0])
            reporter = reporter or meta.get("reporter")
            report_date = report_date or meta.get("date")
        except json.JSONDecodeError:
            pass

    # 从后往前删除指南段，避免索引错位
    removed = 0
    for para in reversed(paragraphs):
        if paragraph_is_guide(para):
            remove_paragraph(para)
            removed += 1
    logs.append(f"已移除 {removed} 段填写说明/指南")

    filled_keys: set[str] = set()

    # 记录模板标题（指南已删除后再扫描，避免段落引用失效）
    heading_slots: list[tuple[Paragraph, str, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        section_key = is_template_heading(para)
        if section_key and section_key not in {"reporter", "date"}:
            heading_slots.append((para, section_key, text))

    for anchor, section_key, heading_text in heading_slots:
        content_parts = sections.get(section_key, [])
        if not content_parts:
            continue

        # 删除标题后紧邻的占位/空段
        nxt = anchor._p.getnext()
        while nxt is not None:
            tag = nxt.tag.split("}")[-1]
            if tag != "p":
                break
            child = Paragraph(nxt, anchor._parent)
            child_text = child.text.strip()
            if is_template_heading(child):
                break
            if child_text and not (paragraph_is_red_instruction(child) or paragraph_is_guide(child)):
                break
            remove_paragraph(child)
            nxt = anchor._p.getnext()

        insert_anchor = anchor
        for part in content_parts:
            new_para = insert_paragraph_after(insert_anchor, part)
            apply_body_style(new_para)
            insert_anchor = new_para

        filled_keys.add(section_key)
        logs.append(f"已填入章节「{heading_text[:40]}」共 {len(content_parts)} 段")

    paragraphs = list(doc.paragraphs)
    for para in paragraphs:
        text = para.text.strip()
        if "汇报人" in text and reporter:
            set_paragraph_text(para, f"                                            汇报人：{reporter}")
            logs.append(f"已更新汇报人: {reporter}")
        elif "日期：" in text and report_date:
            set_paragraph_text(para, f"                                            日期：{report_date}")
            logs.append(f"已更新日期: {report_date}")

    missing = [d["key"] for d in SECTION_DEFS if d["key"] not in filled_keys and not sections.get(d["key"])]
    if missing:
        logs.append(f"警告: 以下内容章节未匹配到模板位置: {', '.join(missing)}")

    empty_sections = [d["key"] for d in SECTION_DEFS if not sections.get(d["key"])]
    if empty_sections:
        logs.append(f"提示: 输入内容缺少章节: {', '.join(empty_sections)}")

    doc.save(str(dst))
    logs.append(f"已生成: {output_path}")
    return logs


def main() -> None:
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "缺少 JSON 参数"}, ensure_ascii=False))
        sys.exit(1)

    try:
        payload = json.loads(sys.argv[1])
    except json.JSONDecodeError as e:
        print(json.dumps({"ok": False, "error": f"JSON 解析失败: {e}"}, ensure_ascii=False))
        sys.exit(1)

    template_path = payload.get("template_path")
    output_path = payload.get("output_path")
    content_path = payload.get("content_path")
    content_text = payload.get("content_text")
    content_kind = payload.get("content_kind", "text")
    reporter = payload.get("reporter")
    report_date = payload.get("report_date")

    if not template_path or not output_path:
        print(json.dumps({"ok": False, "error": "template_path 与 output_path 必填"}, ensure_ascii=False))
        sys.exit(1)

    try:
        sections = load_sections(content_path, content_text, content_kind)
        logs = fill_template(
            template_path,
            output_path,
            sections,
            reporter=reporter,
            report_date=report_date,
        )
        result = {
            "ok": True,
            "results": [{"input": content_path or "(text)", "output": output_path}],
            "logs": logs,
        }
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"ok": False, "error": str(e)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
